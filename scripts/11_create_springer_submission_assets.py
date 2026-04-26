"""Create Springer Nature / Journal of Neural Transmission submission assets.

The generated files are intentionally conservative: they target the
subscription publishing route in a hybrid Springer Nature journal and avoid
claims of cure, clinical recommendation, or definitive prevention.
"""

from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SUB = ROOT / "submission_package"
MAN = ROOT / "manuscript"
OUT = SUB / "springer_nature_jnt_subscription_submission"
OUT.mkdir(parents=True, exist_ok=True)

TITLE = (
    "Artificial intelligence-assisted evidence mapping and computational prioritisation "
    "of preventive and disease-modifying strategies for Parkinson's disease: "
    "a systematic evidence-to-discovery study"
)

GITHUB = "https://github.com/hssling/pd-discovery-benchmark-dashboard"
GITHUB_RELEASE = "https://github.com/hssling/pd-discovery-benchmark-dashboard/releases/tag/v1.0.0"
GITHUB_ZENODO_TRIGGER_RELEASE = "https://github.com/hssling/pd-discovery-benchmark-dashboard/releases/tag/v1.0.1"
HF_EVID = "https://huggingface.co/datasets/hssling/parkinsons-evidence-to-discovery-prioritisation"
KG_EVID = "https://www.kaggle.com/datasets/jkhospital/parkinsons-evidence-to-discovery-prioritisation"
HF_BENCH = "https://huggingface.co/datasets/hssling/pd-discovery-benchmark-dashboard"
KG_BENCH = "https://www.kaggle.com/datasets/jkhospital/pd-discovery-benchmark"


def add_paragraphs(doc: Document, paragraphs: list[str]) -> None:
    for text in paragraphs:
        if text == "":
            doc.add_paragraph("")
        else:
            doc.add_paragraph(text)


def set_style(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)


def save_docx(filename: str, title: str, sections: list[tuple[str, list[str]]]) -> None:
    doc = Document()
    set_style(doc)
    doc.add_heading(title, level=1)
    for heading, paragraphs in sections:
        doc.add_heading(heading, level=2)
        add_paragraphs(doc, paragraphs)
    doc.save(OUT / filename)


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8").strip()


data_availability = read_text(SUB / "data_availability_statement.md")
ethics = read_text(SUB / "ethics_statement.md")
competing = read_text(SUB / "competing_interests.md")

cover_sections = [
    (
        "Correspondence",
        [
            "To: Editors-in-Chief, Journal of Neural Transmission",
            "Submission system: Springer Nature Editorial Manager / Springer Nature submission portal",
            "Manuscript type requested: Review / Research synthesis resource article, depending on editorial classification",
            "Publishing route requested: subscription publishing model, not Open Choice / open access, unless a waiver or institutional agreement is confirmed.",
        ],
    ),
    (
        "Cover Letter",
        [
            "Dear Editors,",
            "",
            f"We submit the manuscript entitled \"{TITLE}\" for consideration in Journal of Neural Transmission.",
            "The manuscript presents an evidence-to-discovery study integrating AI-assisted evidence mapping, clinical-trial mining, public transcriptomic validation, pathway analysis, drug-repurposing prioritisation, and an Indian/public-health prevention framework for Parkinson's disease.",
            "The work is positioned as candidate prioritisation and hypothesis generation. It does not claim that a cure has been found, does not recommend off-label clinical use, and does not assert that any individual intervention definitively prevents Parkinson's disease.",
            "We believe the manuscript fits the journal's translational neuroscience scope because it connects epidemiological prevention evidence, disease-modifying therapeutic targets, omics/pathway recurrence, drug-target resources, and validation-ready experimental models.",
            "All generated data, code, figures, benchmark files, and audit materials are publicly available through the repositories listed in the Data Availability Statement. The GitHub repository includes CI validation and a v1.0.0 release archive suitable for Zenodo DOI minting.",
            "",
            "This manuscript is not under consideration elsewhere. All authors must approve the final submission version before upload.",
            "",
            "Sincerely,",
            "Corresponding author: To be completed by study team",
            "Contact email: To be completed by study team",
        ],
    ),
    (
        "Official Journal/Fee Basis Checked",
        [
            "Journal of Neural Transmission is listed by Springer Nature as a hybrid journal.",
            "Springer Nature states that authors can choose subscription publishing or open access after acceptance; the open-access route carries an APC, while the subscription publishing route has no APC charges.",
            "Official source: https://link.springer.com/journal/702/how-to-publish-with-us",
        ],
    ),
]
save_docx("01_cover_letter_JNT_subscription_route.docx", "Cover Letter", cover_sections)

title_sections = [
    (
        "Title",
        [TITLE],
    ),
    (
        "Author Information",
        [
            "Authors: To be completed by study team.",
            "Affiliations: To be completed by study team.",
            "Corresponding author: To be completed by study team.",
            "Corresponding author email: To be completed by study team.",
            "ORCID IDs: To be completed where available.",
        ],
    ),
    (
        "Keywords",
        [
            "Parkinson disease; disease modification; prevention; GLP-1 receptor agonist; pathway analysis; drug repurposing",
        ],
    ),
    (
        "Running Title",
        ["AI-assisted PD evidence-to-discovery prioritisation"],
    ),
]
save_docx("02_title_page_JNT.docx", "Title Page", title_sections)

declaration_sections = [
    ("Ethics Approval", [ethics]),
    ("Consent To Participate", ["Not applicable; this study used public aggregate literature, trial, database, and de-identified public omics resources."]),
    ("Consent For Publication", ["Not applicable; no identifiable participant-level information is reported."]),
    ("Competing Interests", [competing]),
    ("Funding", ["No specific funding statement has been provided. To be completed by the study team before submission."]),
    ("Author Contributions", ["To be completed by the study team before submission. AI tools and computational scripts assisted evidence mapping, data processing, drafting, and reproducibility packaging; human authors remain responsible for all scientific claims and the final submitted manuscript."]),
    ("Use Of AI Tools", ["AI-assisted workflows were used for evidence mapping, code generation, data extraction support, drafting support, and reproducibility packaging. Outputs were checked against source files, public database records, and audit logs. AI tools are not listed as authors."]),
    ("Clinical Caution", ["The manuscript prioritises candidate intervention points and validation strategies. It does not claim cure, definitive prevention, or clinical efficacy and does not recommend off-label use outside approved indications or research settings."]),
]
save_docx("03_declarations_JNT.docx", "Statements And Declarations", declaration_sections)

das_sections = [
    ("Data Availability Statement", [data_availability]),
    (
        "Repository And Release Links",
        [
            f"Evidence dataset, Hugging Face: {HF_EVID}",
            f"Evidence dataset, Kaggle: {KG_EVID}",
            f"Benchmark GitHub repository: {GITHUB}",
            f"Benchmark GitHub v1.0.0 release: {GITHUB_RELEASE}",
            f"Benchmark GitHub v1.0.1 Zenodo-trigger release: {GITHUB_ZENODO_TRIGGER_RELEASE}",
            f"Benchmark Hugging Face dataset: {HF_BENCH}",
            f"Benchmark Kaggle dataset: {KG_BENCH}",
        ],
    ),
    (
        "Zenodo DOI Placeholder",
        [
            "Zenodo DOI: To be inserted after Zenodo archives the GitHub v1.0.0 release.",
            "Recommended manuscript wording after DOI creation: The versioned benchmark archive is available via Zenodo at DOI: [insert DOI], corresponding to GitHub release v1.0.0.",
        ],
    ),
]
save_docx("04_data_availability_statement_JNT.docx", "Data Availability Statement", das_sections)

checklist_sections = [
    (
        "Target Journal",
        [
            "Journal: Journal of Neural Transmission",
            "Publisher: Springer Nature",
            "Route: subscription publishing model / non-open-access route to avoid APC.",
            "Do not select Open Choice unless APC funding, institutional coverage, or waiver is confirmed.",
        ],
    ),
    (
        "Springer Nature Requirements Checked",
        [
            "Journal publishing model: hybrid.",
            "Accepted source formats include common word-processing formats such as DOCX or LaTeX.",
            "Abstract target: 150-250 words.",
            "Keywords target: 4-6 keywords.",
            "Declarations and competing-interest information should be entered into the submission interface as required.",
            "Official guidelines: https://link.springer.com/journal/702/submission-guidelines",
            "Official fees and publishing-route page: https://link.springer.com/journal/702/how-to-publish-with-us",
        ],
    ),
    (
        "Files To Upload",
        [
            "Main manuscript DOCX: manuscript_for_submission_JNT.docx",
            "Title page: 02_title_page_JNT.docx",
            "Cover letter: 01_cover_letter_JNT_subscription_route.docx",
            "Declarations: 03_declarations_JNT.docx",
            "Data availability statement: 04_data_availability_statement_JNT.docx",
            "Supplementary file: supplementary_file_for_submission.docx",
            "Checklist: STROBE_PRISMA_checklist_for_submission.docx",
            "Figures and tables: use source files from project figures/ and tables/ folders as required by the submission system.",
        ],
    ),
    (
        "Pre-Upload Human Checks",
        [
            "Confirm author names, affiliations, ORCID IDs, and corresponding author email.",
            "Confirm all authors approve submission.",
            "Confirm manuscript is not under consideration elsewhere.",
            "Confirm no copyrighted third-party figure/table content requires permission.",
            "Insert Zenodo DOI if the GitHub release has been archived before submission.",
            "Keep cautious language: prioritised, candidate, suggestive, hypothesis-generating, requires validation.",
        ],
    ),
]
save_docx("05_submission_checklist_JNT_subscription_route.docx", "Submission Checklist", checklist_sections)

zenodo_sections = [
    (
        "Purpose",
        [
            "Create a DOI-backed archive of the GitHub v1.0.0 release for citation in the manuscript and Data Availability Statement.",
        ],
    ),
    (
        "Known Account Context",
        [
            "Zenodo account email provided by user: hssling@yahoo.com",
            "GitHub repository: https://github.com/hssling/pd-discovery-benchmark-dashboard",
            "GitHub release to archive: https://github.com/hssling/pd-discovery-benchmark-dashboard/releases/tag/v1.0.1",
        ],
    ),
    (
        "Zenodo GitHub Integration Steps",
        [
            "1. Log in to Zenodo with the account connected to GitHub.",
            "2. Open the Zenodo GitHub integration page and confirm that the hssling GitHub account is linked.",
            "3. Enable archiving for repository hssling/pd-discovery-benchmark-dashboard.",
            "4. Use release v1.0.1 as the Zenodo-trigger release if v1.0.0 was created before enabling archiving.",
            "5. Wait for Zenodo to archive the release.",
            "6. Open the Zenodo record and copy the version DOI and concept DOI.",
            "7. Replace the DOI placeholder in the Data Availability Statement and repository README if desired.",
        ],
    ),
    (
        "Official Zenodo Note",
        [
            "Zenodo's GitHub workflow archives GitHub releases and then exposes a DOI. Zenodo support states that DOI pre-reservation is not available before using the GitHub integration release workflow.",
            "Official archive-release guide: https://help.zenodo.org/docs/github/archive-software/github-upload/",
            "Official pre-reservation FAQ: https://support.zenodo.org/help/en-gb/24-github-integration/73-can-i-pre-reserved-a-doi-before-a-github-release",
        ],
    ),
    (
        "DOI Placeholder",
        [
            "Version DOI: To be added after Zenodo archive completion.",
            "Concept DOI: To be added after Zenodo archive completion.",
        ],
    ),
]
save_docx("06_zenodo_github_release_doi_checklist.docx", "Zenodo DOI Checklist", zenodo_sections)

source_main = MAN / "manuscript_full_publication_ready.docx"
if not source_main.exists():
    source_main = MAN / "manuscript_full.docx"
shutil.copy2(source_main, OUT / "manuscript_for_submission_JNT.docx")
shutil.copy2(SUB / "supplementary_file.docx", OUT / "supplementary_file_for_submission.docx")
shutil.copy2(SUB / "STROBE_PRISMA_checklist.docx", OUT / "STROBE_PRISMA_checklist_for_submission.docx")

readme = OUT / "README_submission_package.md"
readme.write_text(
    "\n".join(
        [
            "# Springer Nature Submission Package",
            "",
            "Target journal: Journal of Neural Transmission.",
            "Publisher: Springer Nature.",
            "Publishing route: subscription publishing model / non-open-access route to avoid APC.",
            "",
            "## Files",
            "",
            "- `manuscript_for_submission_JNT.docx`",
            "- `01_cover_letter_JNT_subscription_route.docx`",
            "- `02_title_page_JNT.docx`",
            "- `03_declarations_JNT.docx`",
            "- `04_data_availability_statement_JNT.docx`",
            "- `05_submission_checklist_JNT_subscription_route.docx`",
            "- `06_zenodo_github_release_doi_checklist.docx`",
            "- `supplementary_file_for_submission.docx`",
            "- `STROBE_PRISMA_checklist_for_submission.docx`",
            "",
            "## Public Repository Links",
            "",
            f"- GitHub repository: {GITHUB}",
            f"- GitHub v1.0.0 release: {GITHUB_RELEASE}",
            f"- GitHub v1.0.1 Zenodo-trigger release: {GITHUB_ZENODO_TRIGGER_RELEASE}",
            f"- Hugging Face benchmark dataset: {HF_BENCH}",
            f"- Kaggle benchmark dataset: {KG_BENCH}",
            "",
            "## Human Items To Complete",
            "",
            "- author names, affiliations, ORCID IDs, and corresponding author email;",
            "- funding statement;",
            "- author contributions;",
            "- Zenodo DOI after GitHub release archiving;",
            "- final author approval before upload.",
        ]
    ),
    encoding="utf-8",
)

print(f"Created Springer Nature submission assets in {OUT}")
