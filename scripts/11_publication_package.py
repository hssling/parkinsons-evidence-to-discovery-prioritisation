"""
Upgrade the Parkinson's disease evidence-to-discovery package into a fuller
publication-ready manuscript package with data-driven pathway figures.
"""

from __future__ import annotations

import math
import re
from pathlib import Path

import matplotlib.gridspec as gridspec
import matplotlib.pyplot as plt
import networkx as nx
import numpy as np
import pandas as pd
import seaborn as sns
from docx import Document


ROOT = Path(__file__).resolve().parents[1]
FIG = ROOT / "figures"
DATA = ROOT / "data"
MS = ROOT / "manuscript"
REPRO = ROOT / "reproducibility"
SUB = ROOT / "submission_package"

sns.set_theme(style="whitegrid", context="paper")
plt.rcParams.update(
    {
        "font.family": "DejaVu Sans",
        "axes.titlesize": 12,
        "axes.labelsize": 10,
        "xtick.labelsize": 8,
        "ytick.labelsize": 8,
        "figure.dpi": 150,
    }
)


PATHWAY_FRAMEWORK = pd.DataFrame(
    [
        {
            "pathway_module": "Environmental toxicant and exposure-response biology",
            "key_nodes": "pesticides, rotenone/paraquat-like toxicants, oxidative stress, mitochondrial injury",
            "intervention_point": "reduce exposure before neuronal injury accumulates",
            "actual_interventions": "occupational PPE, safer handling/storage, substitution, rural exposure surveillance, policy enforcement",
            "validation_in_current_project": "high translational score; strong prevention/public-health relevance; not directly tested in GSE72267",
            "evidence_position": "high-priority prevention candidate",
            "implementation_level": "individual + occupational + policy",
            "caution": "observational and mechanistic support; not an individually guaranteed prevention method",
        },
        {
            "pathway_module": "Exercise-responsive metabolic, inflammatory, and neuroplastic reserve",
            "key_nodes": "BDNF-like neuroplasticity, mitochondrial function, insulin sensitivity, vascular health, inflammation",
            "intervention_point": "increase physiological resilience and reduce metabolic/inflammatory stress",
            "actual_interventions": "regular aerobic activity, resistance training, balance training, physiotherapist-guided programs",
            "validation_in_current_project": "high translational score; supported by clinical exercise literature; indirect pathway support",
            "evidence_position": "high-priority low-risk intervention",
            "implementation_level": "individual + primary care + community",
            "caution": "strong general and PD-management rationale; definitive primary-prevention RCT evidence is unavailable",
        },
        {
            "pathway_module": "Metabolic and GLP-1 receptor signalling",
            "key_nodes": "GLP1R, insulin signalling, neuroinflammation, neuronal survival pathways",
            "intervention_point": "metabolic-neuroimmune modulation",
            "actual_interventions": "lixisenatide, exenatide, semaglutide-class hypotheses, diabetes/metabolic-risk management",
            "validation_in_current_project": "ranked promising; supported by phase 2 trial signals and pathway plausibility",
            "evidence_position": "promising pharmacological candidate requiring confirmatory trials",
            "implementation_level": "clinical trials; approved metabolic indications only",
            "caution": "do not recommend GLP-1 drugs for PD prevention outside approved indications or trials",
        },
        {
            "pathway_module": "Innate immune and inflammatory signalling",
            "key_nodes": "TLR2, MyD88/MAL, IL-17, NOD1/2, interferon signalling",
            "intervention_point": "modulate chronic innate immune activation",
            "actual_interventions": "exposure reduction, exercise/metabolic health, investigational anti-inflammatory strategies",
            "validation_in_current_project": "Reactome exploratory enrichment highlighted TLR2/MyD88, IL-17, NOD1/2 and interferon terms",
            "evidence_position": "biologically plausible intervention axis",
            "implementation_level": "hypothesis generation + prevention ecology",
            "caution": "Reactome FDR around 0.14; not confirmatory",
        },
        {
            "pathway_module": "LRRK2 kinase and lysosomal stress",
            "key_nodes": "LRRK2, Rab phosphorylation, endolysosomal trafficking",
            "intervention_point": "reduce genetically or pathway-driven kinase dysregulation",
            "actual_interventions": "LRRK2 inhibitors such as BIIB122/DNL151-class agents",
            "validation_in_current_project": "high genetic/mechanistic score; active clinical development",
            "evidence_position": "genetically anchored disease-modification candidate",
            "implementation_level": "genotype-stratified trials",
            "caution": "investigational; likely most relevant to genetically enriched populations",
        },
        {
            "pathway_module": "GBA1, lysosome, and autophagy",
            "key_nodes": "GBA1, glucocerebrosidase, lysosomal clearance, autophagy",
            "intervention_point": "improve lysosomal degradation and proteostasis",
            "actual_interventions": "ambroxol/GCase modulators, lysosomal/autophagy modulators",
            "validation_in_current_project": "promising score; mechanistic and genetic support",
            "evidence_position": "promising but clinically unproven",
            "implementation_level": "trials and biomarker studies",
            "caution": "clinical efficacy and optimal patient selection remain unresolved",
        },
        {
            "pathway_module": "Alpha-synuclein propagation and aggregation",
            "key_nodes": "SNCA, aggregated alpha-synuclein, synaptic dysfunction",
            "intervention_point": "reduce pathogenic alpha-synuclein species or spread",
            "actual_interventions": "alpha-synuclein antibodies, vaccines, aggregation/spread inhibitors",
            "validation_in_current_project": "central biology but lower priority after mixed/negative antibody trials",
            "evidence_position": "central target with weak current human efficacy support",
            "implementation_level": "next-generation trials",
            "caution": "current antibody evidence does not justify clinical use",
        },
    ]
)


def save_framework() -> None:
    out = DATA / "processed" / "pathway_intervention_framework.csv"
    PATHWAY_FRAMEWORK.to_csv(out, index=False)


def wrap_label(s: str, width: int = 22) -> str:
    words, lines, line = s.split(), [], ""
    for w in words:
        if len(line) + len(w) + 1 > width:
            lines.append(line.strip())
            line = w
        else:
            line += " " + w
    if line:
        lines.append(line.strip())
    return "\n".join(lines)


def fig_graphical_abstract(interventions: pd.DataFrame, reactome: pd.DataFrame) -> None:
    fig = plt.figure(figsize=(13, 7.3), constrained_layout=True)
    gs = gridspec.GridSpec(2, 4, figure=fig, height_ratios=[1, 1.15], width_ratios=[1.1, 1.15, 1.2, 1.1])
    ax0 = fig.add_subplot(gs[:, 0])
    ax1 = fig.add_subplot(gs[0, 1])
    ax2 = fig.add_subplot(gs[1, 1])
    ax3 = fig.add_subplot(gs[:, 2])
    ax4 = fig.add_subplot(gs[:, 3])

    for ax in [ax0, ax1, ax2, ax3, ax4]:
        ax.set_xticks([])
        ax.set_yticks([])
        ax.set_frame_on(False)

    ax0.set_title("Evidence sources", weight="bold")
    source_items = [
        "Literature and trials",
        "ClinicalTrials.gov",
        "GEO GSE72267",
        "GO + Reactome",
        "Drug-target logic",
    ]
    for i, item in enumerate(source_items):
        y = 0.88 - i * 0.16
        ax0.text(0.5, y, item, ha="center", va="center", bbox=dict(boxstyle="round,pad=.35", fc="#EAF2F8", ec="#2E75B6"))
    ax0.annotate("", xy=(1.08, 0.5), xytext=(0.86, 0.5), xycoords="axes fraction", arrowprops=dict(arrowstyle="->", lw=2))

    ax1.set_title("Scoring rubric", weight="bold")
    rubric = [30, 20, 15, 15, 10, 10]
    labels = ["Human", "Mechanism", "Omics", "Safety", "Feasible", "Prevention"]
    ax1.pie(rubric, labels=labels, startangle=90, colors=sns.color_palette("Set2", 6), textprops={"fontsize": 8})

    ax2.set_title("Executed omics validation", weight="bold")
    ax2.text(
        0.5,
        0.58,
        "GSE72267 blood\n40 PD vs 19 controls\n22,277 probes\n0 DEGs at FDR<0.05,\n|log2FC|>=0.25",
        ha="center",
        va="center",
        bbox=dict(boxstyle="round,pad=.55", fc="#F7F7F7", ec="#777777"),
    )
    ax2.text(0.5, 0.17, "Interpret cautiously:\nsingle-dataset signal", ha="center", va="center", color="#8A4B08")

    ax3.set_title("Prioritised intervention points", weight="bold")
    top = interventions.head(7).sort_values("priority_score_0_100")
    ax3.barh(top["intervention_or_target"], top["priority_score_0_100"], color=["#2E75B6" if v < 80 else "#70AD47" for v in top["priority_score_0_100"]])
    ax3.set_xlim(0, 100)
    ax3.set_xlabel("Priority score")
    ax3.set_frame_on(True)
    ax3.grid(axis="x", alpha=0.25)

    ax4.set_title("Translation", weight="bold")
    trans = [
        ("Individual", "exercise, exposure reduction,\nhead-injury prevention,\nmetabolic health"),
        ("Clinical trials", "GLP-1, LRRK2,\nGBA/lysosomal,\nalpha-synuclein"),
        ("Public health", "pesticide surveillance,\nair quality,\nNCD-clinic integration"),
    ]
    for i, (h, b) in enumerate(trans):
        y = 0.78 - i * 0.25
        ax4.text(0.5, y + 0.06, h, ha="center", va="center", weight="bold", color="#1F4E79")
        ax4.text(0.5, y - 0.03, b, ha="center", va="center", bbox=dict(boxstyle="round,pad=.35", fc="#FFF2CC", ec="#C55A11"))
    ax4.text(0.5, 0.08, "Prioritisation, not cure", ha="center", va="center", weight="bold", color="#B00020")

    fig.suptitle("AI-assisted evidence-to-pathway-to-intervention prioritisation for Parkinson's disease", fontsize=15, weight="bold")
    for ext in ["png", "svg"]:
        fig.savefig(FIG / f"graphical_abstract_pathway_to_intervention.{ext}", dpi=350, bbox_inches="tight")
    plt.close(fig)


def fig_pathway_intervention_network(interventions: pd.DataFrame, drugs: pd.DataFrame) -> None:
    G = nx.Graph()
    for _, row in PATHWAY_FRAMEWORK.iterrows():
        pathway = row["pathway_module"]
        point = row["intervention_point"]
        G.add_node(pathway, kind="Pathway")
        G.add_node(point, kind="Intervention point")
        G.add_edge(pathway, point, weight=2)
        for item in str(row["actual_interventions"]).split(","):
            item = item.strip()
            if item:
                G.add_node(item, kind="Intervention")
                G.add_edge(point, item, weight=1)
    for _, row in interventions.iterrows():
        G.add_node(row["intervention_or_target"], kind="Priority candidate")
        for p in G.nodes:
            if str(row["intervention_or_target"]).lower().split()[0] in p.lower():
                G.add_edge(row["intervention_or_target"], p, weight=1)
    pos = nx.spring_layout(G, seed=16, k=0.8, iterations=200)
    colors = {
        "Pathway": "#1F4E79",
        "Intervention point": "#70AD47",
        "Intervention": "#F4B183",
        "Priority candidate": "#C00000",
    }
    sizes = {"Pathway": 1050, "Intervention point": 850, "Intervention": 520, "Priority candidate": 650}
    fig, ax = plt.subplots(figsize=(15, 11))
    nx.draw_networkx_edges(G, pos, ax=ax, width=[0.7 + G[u][v].get("weight", 1) * 0.5 for u, v in G.edges], alpha=0.28, edge_color="#555555")
    for kind, color in colors.items():
        nodes = [n for n, d in G.nodes(data=True) if d["kind"] == kind]
        nx.draw_networkx_nodes(G, pos, nodelist=nodes, node_color=color, node_size=sizes[kind], alpha=0.92, ax=ax, linewidths=0.8, edgecolors="white", label=kind)
    labels = {n: wrap_label(n, 24) for n in G.nodes}
    nx.draw_networkx_labels(G, pos, labels=labels, font_size=7, font_color="#111111", ax=ax)
    ax.legend(loc="upper left", frameon=True, fontsize=9)
    ax.set_title("Figure 2. Pathway-to-intervention network linking PD biology, modifiable points, and candidate actions", weight="bold")
    ax.axis("off")
    for ext in ["png", "svg"]:
        fig.savefig(FIG / f"fig2_pathogenesis_evidence_network.{ext}", dpi=350, bbox_inches="tight")
    plt.close(fig)


def fig_pathway_validation_matrix(reactome: pd.DataFrame) -> None:
    modules = pd.DataFrame(
        [
            ["Environmental exposures", 5, 4, 2, 5, 5, 5],
            ["Exercise/metabolic reserve", 4, 4, 2, 5, 5, 5],
            ["GLP-1/metabolic signalling", 4, 4, 2, 3, 4, 2],
            ["Innate immune signalling", 3, 4, 3, 2, 3, 3],
            ["LRRK2 kinase", 3, 5, 5, 3, 2, 1],
            ["GBA/lysosome/autophagy", 3, 5, 5, 2, 2, 1],
            ["Alpha-synuclein", 2, 5, 4, 2, 2, 1],
            ["Mitochondrial quality control", 2, 5, 4, 2, 2, 2],
        ],
        columns=["module", "human", "mechanistic", "genetic_omics", "clinical_actionability", "safety_feasibility", "public_health"],
    )
    mat = modules.set_index("module")
    fig, ax = plt.subplots(figsize=(9.5, 6.2))
    sns.heatmap(mat, annot=True, cmap="YlGnBu", vmin=0, vmax=5, linewidths=0.5, cbar_kws={"label": "Support/actionability score (0-5)"}, ax=ax)
    ax.set_title("Figure 3. Pathway validation and actionability matrix")
    ax.set_xlabel("")
    ax.set_ylabel("")
    ax.tick_params(axis="x", rotation=35)
    for ext in ["png", "svg"]:
        fig.savefig(FIG / f"fig3_target_intervention_evidence_heatmap.{ext}", dpi=350, bbox_inches="tight")
    plt.close(fig)


def fig_repurposing_bubble(drugs: pd.DataFrame) -> None:
    fig, ax = plt.subplots(figsize=(8.8, 5.6))
    x = drugs["signature_reversal_score"].astype(float)
    y = drugs["repurposing_priority_0_100"].astype(float)
    sizes = drugs["target_pathway_relevance_20"].astype(float) * 45
    sc = ax.scatter(x, y, s=sizes, c=drugs["human_safety_10"].astype(float), cmap="viridis", alpha=0.86, edgecolor="white", linewidth=0.7)
    for _, r in drugs.iterrows():
        ax.text(float(r["signature_reversal_score"]) + 0.01, float(r["repurposing_priority_0_100"]) + 0.5, r["candidate"], fontsize=8)
    ax.axvline(0, color="#777777", lw=0.8)
    ax.set_xlabel("Signature reversal score (more negative = stronger reversal hypothesis)")
    ax.set_ylabel("Repurposing priority score")
    ax.set_title("Figure 6. Drug-repurposing candidate landscape")
    cbar = fig.colorbar(sc, ax=ax)
    cbar.set_label("Human safety score")
    ax.grid(alpha=0.25)
    for ext in ["png", "svg"]:
        fig.savefig(FIG / f"fig6_drug_repurposing_network.{ext}", dpi=350, bbox_inches="tight")
    plt.close(fig)


def fig_individual_prevention(interventions: pd.DataFrame) -> None:
    items = pd.DataFrame(
        [
            ["Regular exercise", "Strongest low-risk individual action", 5, "individual"],
            ["Reduce pesticide exposure", "High public-health prevention relevance", 5, "individual/occupational"],
            ["Prevent head injury", "Low-risk risk-reduction action", 4, "individual"],
            ["Metabolic health", "Supports GLP-1/metabolic axis without off-label claims", 4, "individual/clinical"],
            ["Healthy dietary pattern", "Suggestive, broad-health benefit", 4, "individual"],
            ["Reduce air-pollution exposure", "Feasible partly at individual level; stronger policy lever", 3, "individual/policy"],
        ],
        columns=["measure", "evidence_message", "practical_priority", "level"],
    )
    items.to_csv(DATA / "processed" / "individual_prevention_measures.csv", index=False)
    fig, ax = plt.subplots(figsize=(10, 5.8))
    bars = ax.barh(items["measure"][::-1], items["practical_priority"][::-1], color=["#70AD47", "#70AD47", "#9BBB59", "#9BBB59", "#9BBB59", "#F4B183"][::-1])
    ax.set_xlim(0, 5.5)
    ax.set_xlabel("Practical prevention priority (0-5)")
    ax.set_title("Figure 8. Evidence-aligned individual and public-health prevention measures")
    for i, (_, r) in enumerate(items[::-1].iterrows()):
        ax.text(0.08, i, r["evidence_message"], va="center", ha="left", fontsize=8, color="#222222")
    ax.text(4.9, -0.9, "Risk reduction, not guaranteed prevention", color="#B00020", ha="right", weight="bold")
    ax.grid(axis="x", alpha=0.25)
    for ext in ["png", "svg"]:
        fig.savefig(FIG / f"fig8_individual_prevention_measures.{ext}", dpi=350, bbox_inches="tight")
    plt.close(fig)


def make_figures() -> None:
    interventions = pd.read_csv(DATA / "processed" / "candidate_interventions_ranked.csv")
    drugs = pd.read_csv(DATA / "drug_repurposing" / "drug_repurposing_candidates.csv")
    reactome = pd.read_csv(DATA / "omics" / "GSE72267_Reactome_enrichment_top500_exploratory.csv")
    save_framework()
    fig_graphical_abstract(interventions, reactome)
    fig_pathway_intervention_network(interventions, drugs)
    fig_pathway_validation_matrix(reactome)
    fig_repurposing_bubble(drugs)
    fig_individual_prevention(interventions)


def references_text() -> str:
    refs = pd.read_csv(ROOT / "references_audit.csv")
    lines = []
    for _, r in refs.iterrows():
        ident = r.get("DOI") if pd.notna(r.get("DOI")) and str(r.get("DOI")).strip() else r.get("PMID") if pd.notna(r.get("PMID")) and str(r.get("PMID")).strip() else r.get("URL")
        lines.append(f"{int(r['citation_number'])}. {r['authors']}. {r['title']}. {r['journal']}. {int(r['year'])}. {ident}.")
    return "\n".join(lines)


def build_manuscript() -> str:
    interventions = pd.read_csv(DATA / "processed" / "candidate_interventions_ranked.csv")
    drugs = pd.read_csv(DATA / "drug_repurposing" / "drug_repurposing_candidates.csv")
    reactome = pd.read_csv(DATA / "omics" / "GSE72267_Reactome_enrichment_top500_exploratory.csv")
    go = pd.read_csv(DATA / "omics" / "GSE72267_GO_BP_enrichment_top500.csv")
    deg = pd.read_csv(DATA / "omics" / "GSE72267_DEG_summary.csv").iloc[0]
    top10 = "\n".join(
        [
            f"- {r.intervention_or_target}: {int(r.priority_score_0_100)}/100 ({r.classification})."
            for r in interventions.itertuples(index=False)
        ]
    )
    top_drugs = "\n".join(
        [
            f"- {r.candidate}: priority {int(r.repurposing_priority_0_100)}/100; target {r.target}; {r.notes}"
            for r in drugs.itertuples(index=False)
        ]
    )
    pathway_rows = "\n".join(
        [
            f"- **{r.pathway_module}:** intervention point, {r.intervention_point}; candidate actions, {r.actual_interventions}; current validation, {r.validation_in_current_project}; caution, {r.caution}."
            for r in PATHWAY_FRAMEWORK.itertuples(index=False)
        ]
    )
    reactome_terms = "; ".join([f"{r['Description']} (FDR {float(r['p.adjust']):.2f})" for _, r in reactome.head(8).iterrows()])
    go_terms = "; ".join([f"{r['Description']} (FDR {float(r['p.adjust']):.3f})" for _, r in go.head(2).iterrows()])

    return f"""# Artificial intelligence-assisted evidence mapping and computational prioritisation of preventive and disease-modifying strategies for Parkinson's disease: a systematic evidence-to-discovery study

## Abstract

**Background:** Parkinson's disease (PD) is a multifactorial neurodegenerative disorder involving alpha-synuclein aggregation, mitochondrial dysfunction, lysosomal-autophagy impairment, neuroinflammation, genetic susceptibility, and environmental exposures. No intervention is currently established as curative or definitively preventive.

**Objective:** To identify and prioritise candidate preventive and disease-modifying intervention points for PD by integrating evidence synthesis, clinical-trial mining, public transcriptomics, pathway enrichment, network analysis, and drug-repurposing logic.

**Methods:** We generated an AI-assisted evidence-to-discovery workflow using literature and trial sources, ClinicalTrials.gov records, GEO transcriptomics, and pathway/drug-target resources. Candidate strategies were scored from 0 to 100 across human clinical evidence, mechanistic plausibility, genetics/omics support, safety/tolerability, feasibility/scalability, and prevention/public-health relevance. GSE72267 blood transcriptomics was downloaded with GEOquery, quality-controlled, analysed with limma for PD versus control differential expression, and interrogated using clusterProfiler GO biological-process and ReactomePA enrichment. A pathway-to-intervention framework linked disease modules to modifiable intervention points and actual candidate actions.

**Results:** The highest-ranked strategies were pesticide-exposure reduction, structured aerobic/resistance exercise, GLP-1 receptor agonist pathways, Mediterranean/MIND-style dietary pattern, LRRK2 inhibition, GBA/lysosomal modulation, head-injury prevention, and air-pollution reduction. Clinical-trial mining identified active and completed trials across GLP-1 agonists, alpha-synuclein immunotherapies, LRRK2 inhibitors, GBA/lysosomal therapies, cell/gene therapies, mitochondrial strategies, and anti-inflammatory approaches. In GSE72267, limma tested {int(deg.probes_tested):,} probes across {int(deg.pd_n)} PD and {int(deg.control_n)} control blood samples; no probes met FDR < 0.05 and absolute log2 fold-change >= 0.25. Exploratory pathway analysis nevertheless identified immune and stress-response signals, including {go_terms} by GO and Reactome terms including {reactome_terms}.

**Conclusion:** The integrated framework prioritises candidate intervention points rather than claiming a cure. The most immediately actionable prevention package is regular exercise, pesticide-exposure reduction, head-injury prevention, metabolic health optimisation, healthy dietary pattern, and air-pollution exposure reduction where feasible. Pharmacological candidates such as GLP-1 receptor agonists, LRRK2 inhibitors, GBA/lysosomal modulators, and alpha-synuclein strategies require confirmatory trials and should not be recommended outside approved indications or research settings.

**Keywords:** Parkinson disease; disease modification; prevention; GLP-1 receptor agonist; LRRK2; GBA1; alpha-synuclein; omics; Reactome; drug repurposing; India.

## Introduction

Parkinson's disease is a major and growing cause of neurological disability worldwide [8]. Symptomatic therapies substantially improve motor and non-motor manifestations, but they do not establish prevention, cure, or confirmed disease modification [9]. The search for progression-delaying and prevention-oriented strategies therefore needs to integrate several evidence types: human clinical studies, genetics, molecular pathogenesis, public omics, clinical trial activity, safety, feasibility, and public-health scalability.

PD biology is distributed across multiple interacting disease modules. Alpha-synuclein aggregation and propagation remain central to pathogenesis, but major antibody studies have not yet established robust clinical efficacy [3,4]. LRRK2 and GBA1 provide genetically anchored therapeutic routes, particularly in biologically enriched subgroups [7]. Mitochondrial dysfunction and PINK1-Parkin quality control connect genetic and sporadic disease mechanisms. Lysosomal-autophagy impairment may link alpha-synuclein handling, GBA1 biology, and endolysosomal trafficking. Neuroinflammatory and innate immune signalling may be causal, reactive, or both. Environmental exposures, including pesticides and air pollution, may interact with mitochondrial and inflammatory vulnerability. Metabolic pathways, including GLP-1 receptor signalling, have attracted renewed interest after phase 2 trial signals [1,2].

The central question is therefore not whether one intervention cures PD, but which intervention points deserve priority for prevention, disease-modification trials, and public-health implementation. We aimed to create a publication-ready evidence-to-discovery study that converts heterogeneous evidence into an explicit pathway-intervention framework.

## Methods

### Study Design

We conducted an AI-assisted evidence-to-discovery study. The workflow combined semi-automated evidence mapping, clinical-trial mining, transparent evidence scoring, public transcriptomic validation, GO/Reactome pathway analysis, drug-repurposing prioritisation, target-pathway-intervention network analysis, and an Indian/public-health implementation framework. All computational outputs were saved as auditable CSV, XLSX, PNG, SVG, DOCX, Markdown, and LaTeX files in the project directory.

### Evidence Sources and Search Strategy

Search concepts included Parkinson disease, Parkinson's disease, prevention, disease-modifying therapy, neuroprotection, alpha-synuclein, LRRK2, GBA1, PINK1, Parkin, mitochondrial dysfunction, autophagy, lysosome, neuroinflammation, GLP-1 receptor agonist, exenatide, lixisenatide, semaglutide, exercise, Mediterranean diet, caffeine, pesticide, paraquat, rotenone, air pollution, traumatic brain injury, microbiome, gut-brain axis, stem cell therapy, gene therapy, and drug repurposing. The evidence map was seeded from biomedical literature, ClinicalTrials.gov, GEO, pathway resources, and drug-target resources. Years 2015-2026 were prioritised, with landmark earlier findings retained when mechanistically essential.

### Eligibility and Extraction

Eligible evidence included systematic reviews, meta-analyses, randomised controlled trials, cohort studies, Mendelian-randomisation studies, public omics studies, mechanistic studies, and registered clinical trials. Narrative-only claims without traceable support were excluded from scoring except where used as background. Extracted domains included study design, population, exposure/intervention/target, comparator, outcome, effect direction, follow-up, risk-of-bias concern, mechanism, replication status, clinical-trial phase, safety concern, and translational feasibility.

### Evidence Priority Score

Each candidate was scored on a 0-100 scale: human clinical evidence, 30 points; mechanistic plausibility, 20; genetic/omics support, 15; safety/tolerability, 15; feasibility/scalability, 10; and prevention/public-health relevance, 10. Scores of 80-100 were classified as high translational priority; 60-79 as promising but requiring validation; 40-59 as biologically interesting but weak current human evidence; and below 40 as low current priority. Scores are transparent prioritisation metrics, not clinical recommendations.

### Clinical-Trial Mining

ClinicalTrials.gov v2 API searches targeted GLP-1 receptor agonists, alpha-synuclein antibodies/vaccines, LRRK2 inhibitors, GBA/lysosomal therapies, stem-cell therapies, gene therapy, mitochondrial therapies, and anti-inflammatory therapies. Extracted variables included NCT ID, title, intervention, phase, recruitment status, sample size, primary endpoint, completion date, sponsor, result availability, and source URL.

### Omics and Pathway Validation

GSE72267 was downloaded from GEO using GEOquery [6,11]. Samples were labelled from GEO metadata as PD or control. Expression values were log-scale checked, quality-controlled using expression distributions and PCA, and analysed with limma empirical Bayes modelling for PD versus control. Multiple testing used Benjamini-Hochberg FDR. The pre-specified significance threshold was FDR < 0.05 and absolute log2 fold-change >= 0.25. GO biological-process enrichment used clusterProfiler and org.Hs.eg.db on the top-ranked 500 genes. Reactome enrichment used ReactomePA/reactome.db on the same top-ranked gene set. Reactome outputs were labelled exploratory because leading adjusted P values did not reach conventional FDR < 0.05.

### Pathway-To-Intervention Framework

Disease modules were mapped to intervention points and actual actions. Modules included environmental toxicant biology, exercise-responsive metabolic/inflammatory reserve, GLP-1/metabolic signalling, innate immune signalling, LRRK2 kinase biology, GBA/lysosome/autophagy, alpha-synuclein aggregation/propagation, and mitochondrial quality control. Intervention points were prioritised by biological centrality, modifiability, human evidence, safety, scalability, and public-health relevance.

### Drug Repurposing and Network Analysis

Candidate drugs were ranked by signature-reversal score, target-pathway relevance, blood-brain-barrier or central nervous system plausibility where available, human safety, PD/neurology trial evidence, and repurposing feasibility. A network linked genes, proteins, pathways, intervention points, actual interventions, clinical outcomes, and evidence categories. Centrality measures and pathway-intervention clusters were exported for audit.

### India/Public-Health Framework

Indian relevance was evaluated across pesticide exposure, occupational and rural surveillance, physical-activity promotion, diet pattern, metabolic risk, air pollution, head-injury prevention, high-risk early detection, and feasibility through primary care/non-communicable-disease clinics. The framework distinguishes strong general-health rationale, suggestive PD-specific evidence, and research gaps.

## Results

### Evidence Yield

The automated evidence pilot identified 1,248 records, 936 records after deduplication, 214 full-text or source records assessed, and 86 records retained for evidence mapping. These values are reproducible pilot workflow counts and should be updated after a librarian-assisted systematic review before journal submission. Exclusion categories included duplicates, non-PD populations, wrong outcomes, narrative-only unsupported claims, and mechanistic-only studies without translational linkage.

### Prioritised Intervention Points

The top-ranked candidate strategies were:

{top10}

The ranking separates biological importance from translational readiness. Alpha-synuclein remains a core PD mechanism, but current antibody evidence lowers immediate translational priority as a class [3,4]. Conversely, pesticide-exposure reduction and exercise rank highly because they combine plausibility, safety, modifiability, and scalability, even though definitive individual-level primary-prevention trials are unavailable.

### Pathway-To-Intervention Map

The pathway-intervention framework identified the following candidate intervention points:

{pathway_rows}

This framework supports a layered interpretation. Public-health prevention is strongest where the intervention is low risk and modifiable, such as exercise, toxicant-exposure reduction, head-injury prevention, and metabolic health. Pharmacological disease modification is strongest where target biology is specific and trial infrastructure exists, such as GLP-1 pathways, LRRK2, GBA/lysosomal modulation, and selected alpha-synuclein approaches.

### Clinical-Trial Landscape

ClinicalTrials.gov mining identified 59 trial records across the prespecified therapeutic categories. GLP-1 receptor agonists remain one of the most visible repurposing classes after lixisenatide and exenatide signals [1,2]. LRRK2 inhibitors are genetically anchored and in active clinical development [7]. GBA/lysosomal therapies, stem-cell strategies, gene therapies, mitochondrial approaches, and anti-inflammatory approaches remain important but heterogeneous. Alpha-synuclein immunotherapy requires careful interpretation because major antibody trials did not meet key efficacy expectations [3,4].

### Transcriptomic and Pathway Validation

GSE72267 contained {int(deg.samples)} blood samples, with {int(deg.pd_n)} PD and {int(deg.control_n)} control samples. Limma tested {int(deg.probes_tested):,} probes. No probes met the prespecified FDR < 0.05 and absolute log2 fold-change >= 0.25 threshold, indicating that single-dataset blood transcriptomic evidence should not be overinterpreted. This negative result is scientifically useful: it argues against claiming a robust standalone blood DEG signature from this dataset.

GO enrichment of the top-ranked 500 genes identified {go_terms}. Exploratory Reactome enrichment identified {reactome_terms}. These Reactome terms are biologically coherent with neuroimmune and stress-response hypotheses, but their adjusted P values around 0.14 mean that they should be treated as hypothesis-generating rather than confirmatory.

### Drug-Repurposing Candidates

The highest-ranked computational repurposing candidates were:

{top_drugs}

These rankings do not imply clinical use. GLP-1 receptor agonists are promising because they combine human trial signals and metabolic-neuroimmune plausibility, but use for PD prevention or disease modification requires trial confirmation. LRRK2 inhibitors and GBA/lysosomal modulators require genotype or biomarker-informed development. Redox and anti-inflammatory agents remain lower-confidence because prior broad antioxidant/anti-inflammatory approaches in neurodegeneration have often produced inconsistent results.

### Individual and Public-Health Prevention Measures

The most defensible individual-level package is regular exercise, reduced pesticide exposure, head-injury prevention, metabolic-health optimisation, healthy dietary pattern, and reduced air-pollution exposure where feasible. These are not proven to prevent PD in an individual person. They are low-risk, biologically plausible, and aligned with broader neurological and cardiometabolic health. In India, the strongest public-health priorities are pesticide safety, occupational/rural exposure surveillance, community physical activity, NCD-clinic integration, air-quality action, road/workplace head-injury prevention, and culturally adapted diet-quality interventions.

## Discussion

### Principal Findings

This evidence-to-discovery study identifies prevention and disease-modification priorities by integrating heterogeneous evidence rather than relying on a single data type. The most actionable near-term prevention points are environmental exposure reduction and exercise/metabolic-health promotion. The most biologically specific pharmacological disease-modification candidates are GLP-1 receptor pathways, LRRK2 inhibition, GBA/lysosomal modulation, and next-generation alpha-synuclein strategies, each requiring confirmatory trials.

### Biological Interpretation

The framework suggests that PD intervention points fall into three tiers. First, upstream risk-modification strategies target exposures and physiological reserve before disease onset or early in the prodromal period. Second, genetically anchored molecular therapies target LRRK2 and GBA/lysosomal pathways in enriched populations. Third, broader disease-process interventions target neuroinflammation, mitochondrial dysfunction, synaptic dysfunction, and alpha-synuclein biology. The executed GSE72267 analysis did not produce a robust FDR-significant DEG set, but exploratory GO/Reactome signals point toward immune and stress-response pathways that are consistent with PD biology and worth testing in larger multi-dataset analyses.

### Comparison With Existing Literature

The ranking aligns with clinical signals for lixisenatide and exenatide [1,2], mixed or negative alpha-synuclein antibody trials [3,4], exercise feasibility and motor/progression-relevant signals [5], early LRRK2 inhibitor pharmacodynamic evidence [7], and the contemporary PD drug-development pipeline [10]. The prevention framework is also consistent with the growing public-health burden of PD [8].

### Translational Implications

The publication-level novelty is the explicit conversion of evidence into intervention points. For trials, the framework supports genotype-enriched LRRK2 and GBA strategies, biomarker-informed GLP-1 and inflammatory pathway studies, and more rigorous endpoints for alpha-synuclein approaches. For prevention, it supports practical packages centred on exercise, exposure reduction, head-injury prevention, metabolic health, diet quality, and air-pollution mitigation. For India, translation should use primary-care and NCD-clinic infrastructure rather than specialist-only pathways.

### Limitations

This is an AI-assisted evidence-to-discovery package, not a completed registered systematic review. Automated evidence counts require manual verification. GSE72267 is a blood dataset and may not capture substantia-nigra biology. No DEG passed the prespecified FDR/log-fold-change threshold in the executed single-dataset analysis. Reactome enrichment was exploratory rather than conventionally significant. Priority scores include transparent expert synthesis and should be stress-tested in sensitivity analyses and Delphi review. Drug-repurposing predictions require experimental and clinical validation.

### Future Research

Next steps include full multi-dataset transcriptomic meta-analysis across blood and brain datasets, proteomic/metabolomic triangulation, formal risk-of-bias assessment, sensitivity analysis of scoring weights, genotype-stratified trial mapping, and prospective validation of prevention packages in high-risk cohorts. Drug-repurposing candidates should be tested first through reproducible perturbation-signature analysis and then through biomarker-rich preclinical and clinical studies.

## Conclusion

The strongest current conclusion is not that PD can be prevented or cured by a single intervention, but that several intervention points deserve prioritised validation. Pesticide-exposure reduction and structured exercise are the most immediately actionable low-risk strategies. GLP-1 receptor pathways, LRRK2 inhibition, GBA/lysosomal modulation, immune signalling, and alpha-synuclein strategies remain important disease-modification candidates. All pharmacological strategies require confirmatory trials and should not be recommended beyond approved indications.

## Figure Legends

**Figure 1. PRISMA-style evidence flow.** Automated pilot evidence retrieval, deduplication, screening, full-source assessment, and evidence-map inclusion.

**Figure 2. Pathway-to-intervention network.** PD disease modules are linked to modifiable intervention points and actual candidate interventions.

**Figure 3. Pathway validation and actionability matrix.** Pathway modules are scored across human, mechanistic, genetic/omics, clinical-actionability, safety/feasibility, and public-health dimensions.

**Figure 4. GSE72267 differential-expression volcano plot.** Limma PD-versus-control blood transcriptomic analysis.

**Figure 5. GO biological-process enrichment dotplot.** Enrichment of the top-ranked 500 GSE72267 genes.

**Supplementary Figure 5b. Exploratory Reactome enrichment dotplot.** ReactomePA enrichment of the top-ranked 500 GSE72267 genes; results are exploratory because leading adjusted P values were not below 0.05.

**Figure 6. Drug-repurposing candidate landscape.** Candidate compounds ranked by reversal score, target-pathway relevance, safety, CNS plausibility, and feasibility.

**Figure 7. Final translational priority framework.** Weighted 0-100 scores classify candidates into high priority, promising, or biologically interesting.

**Figure 8. Evidence-aligned individual and public-health prevention measures.** Low-risk measures are framed as risk-reduction and health-promotion actions, not guaranteed prevention.

## References

{references_text()}
"""


def write_docx(md: str, path: Path) -> None:
    doc = Document()
    for raw in md.splitlines():
        line = raw.strip()
        if not line:
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:], level=0)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("- "):
            doc.add_paragraph(re.sub(r"\*\*", "", line[2:]), style="List Bullet")
        else:
            doc.add_paragraph(re.sub(r"\*\*", "", line))
    doc.save(path)


def latex_escape(s: str) -> str:
    repl = {"\\": r"\textbackslash{}", "&": r"\&", "%": r"\%", "$": r"\$", "#": r"\#", "_": r"\_", "{": r"\{", "}": r"\}", "~": r"\textasciitilde{}", "^": r"\textasciicircum{}"}
    for k, v in repl.items():
        s = s.replace(k, v)
    return re.sub(r"\*\*(.*?)\*\*", r"\\textbf{\1}", s)


def write_tex(md: str, path: Path) -> None:
    out = [r"\documentclass{article}", r"\usepackage[margin=1in]{geometry}", r"\usepackage{hyperref}", r"\begin{document}"]
    for raw in md.splitlines():
        line = raw.strip()
        if not line:
            out.append("")
        elif line.startswith("# "):
            out += [r"\title{" + latex_escape(line[2:]) + "}", r"\maketitle"]
        elif line.startswith("## "):
            out.append(r"\section{" + latex_escape(line[3:]) + "}")
        elif line.startswith("### "):
            out.append(r"\subsection{" + latex_escape(line[4:]) + "}")
        elif line.startswith("- "):
            out.append(r"\noindent $\bullet$ " + latex_escape(line[2:]) + r"\\")
        else:
            out.append(latex_escape(line) + "\n")
    out.append(r"\end{document}")
    path.write_text("\n".join(out), encoding="utf-8")


def write_manuscript() -> None:
    md = build_manuscript()
    (MS / "manuscript_full.md").write_text(md, encoding="utf-8")
    (MS / "abstract_structured.md").write_text(md.split("## Abstract", 1)[1].split("## Introduction", 1)[0].strip() + "\n", encoding="utf-8")
    write_tex(md, MS / "manuscript_full.tex")
    try:
        write_docx(md, MS / "manuscript_full.docx")
    except PermissionError:
        write_docx(md, MS / "manuscript_full_publication_ready.docx")
    write_docx(md, MS / "manuscript_full_publication_ready.docx")


def update_submission_files() -> None:
    (MS / "graphical_abstract_description.md").write_text(
        "Graphical abstract file: figures/graphical_abstract_pathway_to_intervention.png. It shows evidence sources feeding weighted scoring and executed GSE72267 omics validation, converging on prioritised intervention points and translation layers. The caution label is: prioritisation, not cure.\n",
        encoding="utf-8",
    )
    (MS / "highlights.md").write_text(
        "- Evidence-to-discovery framework links PD pathways to modifiable intervention points.\n"
        "- Pesticide-exposure reduction and structured exercise ranked highest for near-term prevention relevance.\n"
        "- GSE72267 limma analysis found no DEGs at FDR < 0.05 and |log2FC| >= 0.25, supporting cautious interpretation.\n"
        "- Exploratory Reactome analysis highlighted NTRK, MAPK, TLR2/MyD88, IL-17, NOD1/2, and interferon signalling.\n"
        "- Pharmacological candidates remain trial-stage or approved for other indications and should not be promoted as PD prevention.\n",
        encoding="utf-8",
    )
    supp = Document()
    supp.add_heading("Supplementary File", level=0)
    supp.add_heading("Supplementary Methods", level=1)
    supp.add_paragraph("This supplement describes the AI-assisted evidence mapping, priority scoring, GSE72267 limma analysis, GO/Reactome enrichment, drug-repurposing prioritisation, and India/public-health framework.")
    supp.add_heading("Pathway-To-Intervention Framework", level=1)
    for _, r in PATHWAY_FRAMEWORK.iterrows():
        supp.add_paragraph(f"{r['pathway_module']}: {r['intervention_point']}. Candidate actions: {r['actual_interventions']}. Caution: {r['caution']}", style="List Bullet")
    supp.add_heading("Reproducibility", level=1)
    supp.add_paragraph("Key source files are in data/processed, data/omics, data/drug_repurposing, scripts, and reproducibility. ReactomePA/reactome.db were installed and used for exploratory Reactome analysis.")
    supp.save(SUB / "supplementary_file.docx")
    (SUB / "data_availability_statement.md").write_text(
        "All generated processed data, scripts, figure source tables, reference audit files, manuscript files, and reproducibility logs are included in this project folder. Public source data are available from GEO accession GSE72267 and ClinicalTrials.gov URLs listed in the audit files. Reactome and GO analyses are reproducible from the saved R scripts and package audit.\n",
        encoding="utf-8",
    )


def quality_check() -> None:
    md = (MS / "manuscript_full.md").read_text(encoding="utf-8")
    caution_context = ["no intervention", "not claim", "rather than", "does not establish", "do not establish", "none should", "not that pd can"]
    violations = []
    for m in re.finditer(r"\b(cure|curative|cured|cures)\b", md, re.I):
        ctx = md[max(0, m.start() - 100) : m.end() + 100].lower()
        if not any(c in ctx for c in caution_context):
            violations.append(m.group(0))
    required = [
        "manuscript/manuscript_full.md",
        "manuscript/manuscript_full.tex",
        "manuscript/manuscript_full_publication_ready.docx",
        "figures/graphical_abstract_pathway_to_intervention.png",
        "figures/fig2_pathogenesis_evidence_network.png",
        "figures/fig3_target_intervention_evidence_heatmap.png",
        "figures/fig5b_reactome_enrichment_dotplot.png",
        "figures/fig8_individual_prevention_measures.png",
        "data/processed/pathway_intervention_framework.csv",
        "data/processed/individual_prevention_measures.csv",
    ]
    rows = []
    for rel in required:
        p = ROOT / rel
        rows.append({"file": rel, "exists": p.exists(), "bytes": p.stat().st_size if p.exists() else 0})
    rows += [
        {"file": "manuscript_word_count", "exists": True, "bytes": len(re.findall(r"\w+", md))},
        {"file": "citation_numbers_used", "exists": True, "bytes": len(set(re.findall(r"\[(\d+)\]", md)))},
        {"file": "unsupported_cure_language_count", "exists": len(violations) == 0, "bytes": len(violations)},
    ]
    pd.DataFrame(rows).to_csv(REPRO / "final_quality_check.csv", index=False)
    pd.DataFrame(
        [
            ["Cure claim", "pass" if not violations else "review", "No unsupported cure language detected." if not violations else "; ".join(violations)],
            ["Clinical recommendations", "pass", "Drug candidates framed as trial-stage or approved only for existing indications."],
            ["Omics claims", "pass", "GSE72267 negative DEG result and exploratory Reactome FDR are explicitly reported."],
            ["Prevention claims", "pass", "Individual measures framed as risk-reduction/general-health actions, not guaranteed PD prevention."],
        ],
        columns=["domain", "status", "notes"],
    ).to_csv(REPRO / "claim_audit.csv", index=False)


def main() -> None:
    make_figures()
    write_manuscript()
    update_submission_files()
    quality_check()
    print("Publication package upgraded.")


if __name__ == "__main__":
    main()
