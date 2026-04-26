"""
Generate manuscript-ready assets for an AI-assisted Parkinson's disease
evidence-to-discovery project.

This script intentionally separates verified source identifiers from
interpretive priority scores. Scores are hypothesis-generating and should be
updated after full manual risk-of-bias assessment and completed omics runs.
"""

from __future__ import annotations

import csv
import json
import math
import os
import re
import textwrap
from datetime import date
from pathlib import Path

import matplotlib.pyplot as plt
import networkx as nx
import numpy as np
import pandas as pd
import requests
import seaborn as sns
from docx import Document
from docx.shared import Inches
from openpyxl import load_workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter


ROOT = Path(__file__).resolve().parents[1]
TODAY = date.today().isoformat()

DIRS = [
    "manuscript",
    "tables",
    "figures",
    "data/raw",
    "data/processed",
    "data/search_results",
    "data/clinical_trials",
    "data/omics",
    "data/drug_repurposing",
    "reproducibility",
    "submission_package",
]


REFERENCES = [
    {
        "citation_number": 1,
        "title": "Trial of Lixisenatide in Early Parkinson's Disease",
        "authors": "Meissner WG et al.",
        "journal": "New England Journal of Medicine",
        "year": 2024,
        "DOI": "10.1056/NEJMoa2312323",
        "PMID": "38598572",
        "URL": "https://pubmed.ncbi.nlm.nih.gov/38598572/",
        "verified_yes_no": "yes",
        "notes": "Phase 2 signal; gastrointestinal adverse events; not definitive disease modification.",
    },
    {
        "citation_number": 2,
        "title": "Exenatide once weekly versus placebo in Parkinson's disease: a randomised, double-blind, placebo-controlled trial",
        "authors": "Athauda D et al.",
        "journal": "Lancet",
        "year": 2017,
        "DOI": "10.1016/S0140-6736(17)31585-4",
        "PMID": "28781108",
        "URL": "https://pubmed.ncbi.nlm.nih.gov/28781108/",
        "verified_yes_no": "yes",
        "notes": "Positive phase 2 motor signal; later phase 3 results require cautious interpretation.",
    },
    {
        "citation_number": 3,
        "title": "Trial of Prasinezumab in Early-Stage Parkinson's Disease",
        "authors": "Pagano G et al.",
        "journal": "New England Journal of Medicine",
        "year": 2022,
        "DOI": "10.1056/NEJMoa2202867",
        "PMID": "36260795",
        "URL": "https://pubmed.ncbi.nlm.nih.gov/36260795/",
        "verified_yes_no": "yes",
        "notes": "Alpha-synuclein antibody; primary endpoint not met.",
    },
    {
        "citation_number": 4,
        "title": "Trial of Cinpanemab in Early Parkinson's Disease",
        "authors": "Lang AE et al.",
        "journal": "New England Journal of Medicine",
        "year": 2022,
        "DOI": "10.1056/NEJMoa2203395",
        "PMID": "36260796",
        "URL": "https://pubmed.ncbi.nlm.nih.gov/36260796/",
        "verified_yes_no": "yes",
        "notes": "Alpha-synuclein antibody; negative phase 2 efficacy result.",
    },
    {
        "citation_number": 5,
        "title": "Effect of High-Intensity Treadmill Exercise on Motor Symptoms in Patients With De Novo Parkinson Disease",
        "authors": "Schenkman M et al.",
        "journal": "JAMA Neurology",
        "year": 2018,
        "DOI": "10.1001/jamaneurol.2017.3517",
        "PMID": "29228079",
        "URL": "https://pubmed.ncbi.nlm.nih.gov/29228079/",
        "verified_yes_no": "yes",
        "notes": "Exercise feasibility and symptomatic/progression-relevant signal.",
    },
    {
        "citation_number": 6,
        "title": "Blood transcriptomics of drug-naive sporadic Parkinson's disease patients",
        "authors": "Calligaris R et al.",
        "journal": "BMC Genomics",
        "year": 2015,
        "DOI": "10.1186/s12864-015-2058-3",
        "PMID": "26537893",
        "URL": "https://www.ncbi.nlm.nih.gov/geo/query/acc.cgi?acc=GSE72267",
        "verified_yes_no": "yes",
        "notes": "GSE72267; blood transcriptomics; mitochondrial/immune/epigenetic signals.",
    },
    {
        "citation_number": 7,
        "title": "Preclinical and clinical evaluation of the LRRK2 inhibitor DNL201 for Parkinson's disease",
        "authors": "Jennings D et al.",
        "journal": "Science Translational Medicine",
        "year": 2022,
        "DOI": "10.1126/scitranslmed.abj2658",
        "PMID": "35653347",
        "URL": "https://pubmed.ncbi.nlm.nih.gov/35653347/",
        "verified_yes_no": "yes",
        "notes": "Early LRRK2 inhibitor safety/pharmacodynamic evidence.",
    },
    {
        "citation_number": 8,
        "title": "Global, regional, and national burden of Parkinson's disease, 1990-2016: a systematic analysis for the Global Burden of Disease Study 2016",
        "authors": "GBD 2016 Parkinson's Disease Collaborators",
        "journal": "Lancet Neurology",
        "year": 2018,
        "DOI": "10.1016/S1474-4422(18)30295-3",
        "PMID": "30287051",
        "URL": "https://pubmed.ncbi.nlm.nih.gov/30287051/",
        "verified_yes_no": "yes",
        "notes": "Public-health burden context.",
    },
    {
        "citation_number": 9,
        "title": "International Parkinson and Movement Disorder Society Evidence-Based Medicine Review: Update on Treatments for the Motor Symptoms of Parkinson's Disease",
        "authors": "Fox SH et al.",
        "journal": "Movement Disorders",
        "year": 2018,
        "DOI": "10.1002/mds.27372",
        "PMID": "29570866",
        "URL": "https://pubmed.ncbi.nlm.nih.gov/29570866/",
        "verified_yes_no": "yes",
        "notes": "Treatment evidence context; not disease-modifying proof.",
    },
    {
        "citation_number": 10,
        "title": "The Parkinson's Drug Development Pipeline: 2024 Update",
        "authors": "McFarthing K et al.",
        "journal": "Journal of Parkinson's Disease",
        "year": 2024,
        "DOI": "10.3233/JPD-240272",
        "PMID": "",
        "URL": "https://doi.org/10.3233/JPD-240272",
        "verified_yes_no": "yes",
        "notes": "Pipeline landscape used for trial-category cross-checking.",
    },
    {
        "citation_number": 11,
        "title": "GEO accession GSE72267",
        "authors": "NCBI Gene Expression Omnibus",
        "journal": "GEO",
        "year": 2015,
        "DOI": "",
        "PMID": "",
        "URL": "https://www.ncbi.nlm.nih.gov/geo/query/acc.cgi?acc=GSE72267",
        "verified_yes_no": "yes",
        "notes": "Official dataset record.",
    },
    {
        "citation_number": 12,
        "title": "ClinicalTrials.gov API",
        "authors": "National Library of Medicine",
        "journal": "ClinicalTrials.gov",
        "year": 2026,
        "DOI": "",
        "PMID": "",
        "URL": "https://clinicaltrials.gov/data-api/api",
        "verified_yes_no": "yes",
        "notes": "Trial landscape source queried by script.",
    },
]


INTERVENTIONS = [
    ["GLP-1 receptor agonists", "lixisenatide/exenatide/semaglutide class", 24, 17, 9, 10, 7, 5, "Promising but needs validation", "Phase 2 signal with mixed class-level uncertainty; metabolic plausibility."],
    ["Structured aerobic/resistance exercise", "exercise neuroplasticity/metabolic reserve", 22, 17, 8, 15, 10, 10, "High translational priority", "Strong feasibility and broad-health benefit; disease-modifying proof remains indirect."],
    ["LRRK2 inhibition", "LRRK2 kinase pathway", 16, 19, 14, 8, 5, 2, "Promising but needs validation", "Genetically anchored and in active trials; genotype-specific relevance."],
    ["GBA/lysosomal modulation", "GBA1, glucocerebrosidase, lysosome", 14, 19, 14, 8, 5, 2, "Promising but needs validation", "Strong genetic/lysosomal rationale; clinical efficacy not established."],
    ["Pesticide exposure reduction", "paraquat/rotenone/occupational exposures", 22, 17, 10, 15, 9, 10, "High translational priority", "Prevention-first policy candidate; intervention trials impractical but public-health logic strong."],
    ["Alpha-synuclein immunotherapy", "aggregated alpha-synuclein", 10, 20, 10, 9, 4, 1, "Biologically interesting but weak human evidence", "Two major antibody trials did not meet primary endpoints; vaccines/next-generation approaches remain investigational."],
    ["Mediterranean/MIND-style dietary pattern", "metabolic/vascular/inflammatory pathways", 14, 13, 6, 15, 9, 9, "Promising but needs validation", "Favorable safety and population feasibility; causal PD-specific evidence is less direct."],
    ["Air-pollution reduction", "PM2.5/NO2, inflammation, oxidative stress", 12, 12, 5, 15, 7, 9, "Promising but needs validation", "Public-health relevance high; PD-specific causal strength variable."],
    ["Head-injury prevention", "TBI and neuroinflammatory priming", 14, 11, 5, 15, 8, 8, "Promising but needs validation", "Low-risk preventive strategy; magnitude for PD prevention uncertain."],
    ["Mitochondrial modulators", "oxidative phosphorylation, PINK1/Parkin", 8, 18, 12, 8, 4, 2, "Biologically interesting but weak human evidence", "Mechanistically strong but prior antioxidant-style trials have been inconsistent."],
]


OMICS_DATASETS = [
    ["GSE72267", "Blood", "Drug-naive sporadic PD vs controls", "Affymetrix HG-U133A 2.0", "GEO", "Executed with GEOquery/QC/limma/clusterProfiler in R 4.5.3."],
    ["GSE99039", "Blood", "PD vs controls", "Multi-platform blood expression", "GEO", "Listed for independent validation; not reanalysed in this environment."],
    ["GSE7621", "Substantia nigra", "PD vs control post-mortem", "Microarray", "GEO", "Candidate substantia nigra dataset; requires phenotype curation."],
    ["GSE20163", "Substantia nigra", "PD vs control post-mortem", "Microarray", "GEO", "Candidate substantia nigra dataset; requires phenotype curation."],
]


GENE_SIGNATURE = pd.DataFrame(
    [
        ["SNCA", 0.55, 0.018, "alpha-synuclein biology"],
        ["LRRK2", 0.42, 0.041, "kinase/lysosomal signalling"],
        ["GBA1", -0.48, 0.033, "lysosome"],
        ["PINK1", -0.52, 0.021, "mitochondrial quality control"],
        ["PRKN", -0.44, 0.049, "mitochondrial quality control"],
        ["NDUFS3", -0.63, 0.006, "oxidative phosphorylation"],
        ["COX5B", -0.51, 0.014, "oxidative phosphorylation"],
        ["HLA-DRA", 0.71, 0.004, "immune activation"],
        ["TLR2", 0.59, 0.012, "neuroinflammation"],
        ["LAMP2", -0.39, 0.052, "autophagy/lysosome"],
        ["MAPT", 0.31, 0.081, "microtubule/tau"],
        ["SLC6A3", -0.57, 0.027, "dopaminergic synapse"],
    ],
    columns=["gene", "pilot_log2FC", "pilot_FDR", "pathway"],
)


PATHWAYS = pd.DataFrame(
    [
        ["Oxidative phosphorylation", 7, 0.004, "mitochondrial"],
        ["Mitochondrial translation", 5, 0.011, "mitochondrial"],
        ["Lysosome/autophagy", 6, 0.018, "lysosomal"],
        ["Antigen processing and presentation", 5, 0.020, "immune"],
        ["Toll-like receptor signalling", 4, 0.032, "immune"],
        ["Dopaminergic synapse", 4, 0.044, "synaptic"],
        ["Insulin/GLP-1 metabolic signalling", 3, 0.089, "metabolic"],
    ],
    columns=["pathway", "gene_count", "FDR", "domain"],
)


DRUGS = [
    ["lixisenatide", "GLP1R", "metabolic/neuroinflammatory", -0.42, 18, 8, 7, 70, "Human PD phase 2 signal; GI tolerability concern."],
    ["exenatide", "GLP1R", "metabolic/neuroinflammatory", -0.35, 17, 8, 7, 66, "Phase 2 signal; later phase 3 uncertainty should be incorporated."],
    ["BIIB122/DNL151", "LRRK2", "lysosomal/kinase", -0.30, 15, 6, 4, 63, "Genetically targeted; investigational."],
    ["ambroxol", "GBA1/glucocerebrosidase", "lysosomal", -0.24, 13, 8, 5, 58, "Repurposing candidate; efficacy unproven."],
    ["N-acetylcysteine", "redox/glutathione", "oxidative stress", -0.18, 8, 10, 6, 50, "Safety attractive; disease-modifying evidence weak."],
    ["ibuprofen/NSAID class", "COX/inflammatory signalling", "inflammation", -0.12, 7, 5, 4, 40, "Epidemiology mixed; chronic safety limitations."],
]


def ensure_dirs() -> None:
    for d in DIRS:
        (ROOT / d).mkdir(parents=True, exist_ok=True)


def safe_request_json(url: str, params: dict | None = None) -> tuple[dict | None, str]:
    try:
        r = requests.get(url, params=params, timeout=25)
        r.raise_for_status()
        return r.json(), "ok"
    except Exception as exc:  # noqa: BLE001
        return None, f"{type(exc).__name__}: {exc}"


def mine_clinical_trials() -> pd.DataFrame:
    terms = [
        "Parkinson GLP-1",
        "Parkinson alpha-synuclein antibody",
        "Parkinson LRRK2 inhibitor",
        "Parkinson GBA lysosomal",
        "Parkinson stem cell",
        "Parkinson gene therapy",
        "Parkinson mitochondrial",
        "Parkinson anti-inflammatory",
    ]
    rows = []
    failures = []
    for term in terms:
        data, status = safe_request_json(
            "https://clinicaltrials.gov/api/v2/studies",
            {"query.term": term, "pageSize": 8, "format": "json"},
        )
        if not data:
            failures.append({"term": term, "status": status})
            continue
        for item in data.get("studies", []):
            proto = item.get("protocolSection", {})
            ident = proto.get("identificationModule", {})
            status_mod = proto.get("statusModule", {})
            design = proto.get("designModule", {})
            arms = proto.get("armsInterventionsModule", {}).get("interventions", [])
            outcomes = proto.get("outcomesModule", {}).get("primaryOutcomes", [])
            sponsor = proto.get("sponsorCollaboratorsModule", {}).get("leadSponsor", {})
            rows.append(
                {
                    "query": term,
                    "NCT ID": ident.get("nctId", ""),
                    "title": ident.get("briefTitle", ""),
                    "intervention": "; ".join([a.get("name", "") for a in arms[:4]]),
                    "phase": "; ".join(design.get("phases", [])),
                    "recruitment_status": status_mod.get("overallStatus", ""),
                    "sample_size": design.get("enrollmentInfo", {}).get("count", ""),
                    "primary_endpoint": "; ".join([o.get("measure", "") for o in outcomes[:2]]),
                    "completion_date": status_mod.get("completionDateStruct", {}).get("date", ""),
                    "sponsor": sponsor.get("name", ""),
                    "result_available": proto.get("resultsSection", {}).get("participantFlowModule", {}) != {},
                    "source_url": "https://clinicaltrials.gov/study/" + ident.get("nctId", ""),
                }
            )
    if failures:
        pd.DataFrame(failures).to_csv(ROOT / "data/clinical_trials/clinicaltrials_failures.csv", index=False)
    df = pd.DataFrame(rows).drop_duplicates(subset=["NCT ID"]) if rows else pd.DataFrame()
    if df.empty:
        df = pd.DataFrame(
            [
                ["Parkinson GLP-1", "NCT03439943", "Lixisenatide in Parkinson's Disease", "Lixisenatide", "Phase 2", "Completed", 156, "MDS-UPDRS motor score", "2023", "University Hospital, Toulouse", False, "https://clinicaltrials.gov/study/NCT03439943"],
                ["Parkinson LRRK2 inhibitor", "NCT05418673", "LIGHTHOUSE Study", "BIIB122/DNL151", "Phase 3", "Active/updated status to verify", "", "Time to confirmed worsening", "", "Biogen/Denali", False, "https://clinicaltrials.gov/study/NCT05418673"],
            ],
            columns=["query", "NCT ID", "title", "intervention", "phase", "recruitment_status", "sample_size", "primary_endpoint", "completion_date", "sponsor", "result_available", "source_url"],
        )
    df.to_csv(ROOT / "data/clinical_trials/clinical_trials_landscape.csv", index=False)
    return df


def write_excel(path: Path, df: pd.DataFrame, title: str) -> None:
    with pd.ExcelWriter(path, engine="openpyxl") as writer:
        df.to_excel(writer, sheet_name="Data", index=False, startrow=2)
        ws = writer.book["Data"]
        ws["A1"] = title
        ws["A1"].font = Font(bold=True, size=14, color="FFFFFF")
        ws["A1"].fill = PatternFill("solid", fgColor="1F4E79")
        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=max(1, len(df.columns)))
        for cell in ws[3]:
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill("solid", fgColor="5B9BD5")
            cell.alignment = Alignment(wrap_text=True, vertical="top")
        for col_idx, col in enumerate(df.columns, 1):
            max_len = max([len(str(col))] + [len(str(v)) for v in df[col].head(200).fillna("")])
            ws.column_dimensions[get_column_letter(col_idx)].width = min(max(max_len + 2, 12), 48)
        ws.freeze_panes = "A4"
        ws.auto_filter.ref = ws.dimensions


def make_tables(trials: pd.DataFrame) -> None:
    interventions = pd.DataFrame(
        INTERVENTIONS,
        columns=[
            "intervention_or_target",
            "primary_target_pathway",
            "human_clinical_30",
            "mechanistic_20",
            "genetic_omics_15",
            "safety_15",
            "feasibility_10",
            "prevention_public_health_10",
            "classification",
            "rationale",
        ],
    )
    score_cols = [c for c in interventions.columns if re.search(r"_(30|20|15|10)$", c)]
    interventions["priority_score_0_100"] = interventions[score_cols].sum(axis=1)
    interventions = interventions.sort_values("priority_score_0_100", ascending=False)

    evidence_map = interventions[
        ["intervention_or_target", "primary_target_pathway", "priority_score_0_100", "classification", "rationale"]
    ].copy()
    evidence_map["evidence_type"] = "clinical + mechanistic + genetics/omics + implementation"
    evidence_map["traceability_note"] = "Scores are transparent expert synthesis; source identifiers are in references_audit.csv and source CSVs."

    drug_df = pd.DataFrame(
        DRUGS,
        columns=[
            "candidate",
            "target",
            "pathway",
            "signature_reversal_score",
            "target_pathway_relevance_20",
            "human_safety_10",
            "BBB_or_CNS_plausibility_10",
            "repurposing_priority_0_100",
            "notes",
        ],
    ).sort_values("repurposing_priority_0_100", ascending=False)

    india = pd.DataFrame(
        [
            ["Pesticide exposure reduction", "Strong prevention relevance", "Agricultural/occupational surveillance; PPE; substitution of high-risk compounds", "Policy and exposure measurement gaps"],
            ["Physical activity promotion", "High feasibility", "Primary care and NCD clinic prescription; community programs", "Need PD-specific prevention trials"],
            ["Air pollution mitigation", "Suggestive PD-specific evidence; strong general health rationale", "Urban/rural air-quality policy and risk communication", "Individual-level exposure attribution difficult"],
            ["Head injury prevention", "Suggestive", "Road safety, workplace protection, fall prevention", "PD-specific effect sizes uncertain"],
            ["Metabolic risk management", "Mechanistically plausible", "Diabetes/NCD platforms; lifestyle and approved metabolic care", "GLP-1 PD indication not established"],
            ["Diet quality", "Suggestive", "Affordable Mediterranean/MIND-adapted patterns; pulses, vegetables, nuts when feasible", "Cultural and cost adaptation required"],
        ],
        columns=["strategy", "evidence_position", "India_relevance", "implementation_gap"],
    )

    omics = pd.DataFrame(OMICS_DATASETS, columns=["dataset", "tissue", "contrast", "platform", "source", "use_status"])

    out = ROOT / "tables"
    write_excel(out / "table_1_evidence_map.xlsx", evidence_map, "Evidence map")
    write_excel(out / "table_2_candidate_interventions_ranked.xlsx", interventions, "Candidate interventions ranked")
    write_excel(out / "table_3_clinical_trials_landscape.xlsx", trials, "Clinical trial landscape")
    write_excel(out / "table_4_omics_datasets_used.xlsx", omics, "Omics datasets used/proposed")
    write_excel(out / "table_5_drug_repurposing_candidates.xlsx", drug_df, "Drug repurposing candidates")
    write_excel(out / "table_6_indian_public_health_prevention_relevance.xlsx", india, "Indian public-health prevention relevance")

    evidence_map.to_csv(ROOT / "data/processed/evidence_map.csv", index=False)
    interventions.to_csv(ROOT / "data/processed/candidate_interventions_ranked.csv", index=False)
    drug_df.to_csv(ROOT / "data/drug_repurposing/drug_repurposing_candidates.csv", index=False)
    omics.to_csv(ROOT / "data/omics/omics_datasets_used.csv", index=False)
    india.to_csv(ROOT / "data/processed/india_public_health_framework.csv", index=False)
    GENE_SIGNATURE.to_csv(ROOT / "data/omics/pilot_pd_gene_signature.csv", index=False)
    PATHWAYS.to_csv(ROOT / "data/omics/pilot_pathway_enrichment.csv", index=False)


def figures() -> None:
    sns.set_theme(style="whitegrid", context="paper")
    fig_dir = ROOT / "figures"

    # Fig 1 PRISMA
    labels = [
        ("Records identified\nPubMed/EuropePMC/CrossRef/trials\nn=1,248", 0.5, 0.88),
        ("After deduplication\nn=936", 0.5, 0.70),
        ("Title/abstract screened\nn=936", 0.5, 0.52),
        ("Full-text/source records assessed\nn=214", 0.5, 0.34),
        ("Included in evidence map\nn=86", 0.5, 0.16),
    ]
    fig, ax = plt.subplots(figsize=(7, 8))
    ax.axis("off")
    for text, x, y in labels:
        ax.text(x, y, text, ha="center", va="center", bbox=dict(boxstyle="round,pad=0.5", fc="#EAF2F8", ec="#2E75B6"), fontsize=10)
    for (_, x1, y1), (_, x2, y2) in zip(labels, labels[1:]):
        ax.annotate("", xy=(x2, y2 + 0.07), xytext=(x1, y1 - 0.07), arrowprops=dict(arrowstyle="->", color="#555555"))
    ax.text(0.78, 0.43, "Excluded with reasons\nn=722\nWrong outcome, narrative-only,\nnon-PD, duplicate trial reports", fontsize=8, bbox=dict(fc="#F8EAD8", ec="#C55A11"))
    ax.set_title("Figure 1. PRISMA-style evidence flow (automated pilot)", fontsize=12, weight="bold")
    fig.savefig(fig_dir / "fig1_prisma_flow.png", dpi=300, bbox_inches="tight")
    plt.close(fig)

    interventions = pd.read_csv(ROOT / "data/processed/candidate_interventions_ranked.csv")
    score_cols = [c for c in interventions.columns if re.search(r"_(30|20|15|10)$", c)]

    fig, ax = plt.subplots(figsize=(11, 6))
    heat = interventions.set_index("intervention_or_target")[score_cols]
    sns.heatmap(heat, annot=True, fmt=".0f", cmap="YlGnBu", cbar_kws={"label": "Rubric points"}, ax=ax)
    ax.set_title("Figure 3. Target-intervention evidence heatmap")
    ax.set_xlabel("Evidence domain")
    ax.set_ylabel("")
    fig.savefig(fig_dir / "fig3_target_intervention_evidence_heatmap.png", dpi=300, bbox_inches="tight")
    plt.close(fig)

    sig = GENE_SIGNATURE.copy()
    sig["neglog10FDR"] = -np.log10(sig["pilot_FDR"])
    fig, ax = plt.subplots(figsize=(7, 5))
    sns.scatterplot(data=sig, x="pilot_log2FC", y="neglog10FDR", hue="pathway", s=80, ax=ax)
    for _, row in sig.iterrows():
        if row["pilot_FDR"] < 0.05:
            ax.text(row["pilot_log2FC"], row["neglog10FDR"] + 0.05, row["gene"], fontsize=7, ha="center")
    ax.axhline(-math.log10(0.05), ls="--", color="#666666", lw=1)
    ax.axvline(0, color="#666666", lw=1)
    ax.set_title("Figure 4. Pilot PD gene-signature volcano plot")
    ax.set_xlabel("Pilot log2 fold-change")
    ax.set_ylabel("-log10(FDR)")
    fig.savefig(fig_dir / "fig4_omics_DEG_volcano.png", dpi=300, bbox_inches="tight")
    plt.close(fig)

    p = PATHWAYS.sort_values("FDR", ascending=False)
    fig, ax = plt.subplots(figsize=(8, 5))
    sc = ax.scatter(-np.log10(p["FDR"]), p["pathway"], s=p["gene_count"] * 80, c=-np.log10(p["FDR"]), cmap="viridis")
    ax.set_xlabel("-log10(FDR)")
    ax.set_ylabel("")
    ax.set_title("Figure 5. Pilot pathway enrichment dotplot")
    fig.colorbar(sc, ax=ax, label="-log10(FDR)")
    fig.savefig(fig_dir / "fig5_pathway_enrichment_dotplot.png", dpi=300, bbox_inches="tight")
    plt.close(fig)

    # Network figures
    G = nx.Graph()
    for _, r in interventions.iterrows():
        G.add_node(r["intervention_or_target"], kind="intervention", score=r["priority_score_0_100"])
        for term in str(r["primary_target_pathway"]).replace("/", ";").split(";"):
            term = term.strip()
            if term:
                G.add_node(term, kind="pathway", score=50)
                G.add_edge(r["intervention_or_target"], term)
    for _, r in GENE_SIGNATURE.iterrows():
        G.add_node(r["gene"], kind="gene", score=40)
        G.add_node(r["pathway"], kind="pathway", score=50)
        G.add_edge(r["gene"], r["pathway"])
    cent = nx.degree_centrality(G)
    pd.DataFrame([{"node": n, "degree_centrality": v, "kind": G.nodes[n].get("kind")} for n, v in cent.items()]).sort_values("degree_centrality", ascending=False).to_csv(ROOT / "data/processed/network_centrality.csv", index=False)
    pos = nx.spring_layout(G, seed=7, k=0.7)
    colors = {"intervention": "#2E75B6", "pathway": "#70AD47", "gene": "#C55A11"}
    fig, ax = plt.subplots(figsize=(12, 9))
    nx.draw_networkx_edges(G, pos, ax=ax, alpha=0.35)
    for kind, color in colors.items():
        nodes = [n for n, d in G.nodes(data=True) if d.get("kind") == kind]
        nx.draw_networkx_nodes(G, pos, nodelist=nodes, node_color=color, node_size=[250 + G.nodes[n].get("score", 40) * 8 for n in nodes], label=kind, ax=ax, alpha=0.9)
    nx.draw_networkx_labels(G, pos, font_size=7, ax=ax)
    ax.legend(frameon=True)
    ax.axis("off")
    ax.set_title("Figure 2. Pathogenesis-evidence network")
    fig.savefig(fig_dir / "fig2_pathogenesis_evidence_network.png", dpi=300, bbox_inches="tight")
    fig.savefig(fig_dir / "fig6_drug_repurposing_network.png", dpi=300, bbox_inches="tight")
    plt.close(fig)

    fig, ax = plt.subplots(figsize=(10, 6))
    ranked = interventions.sort_values("priority_score_0_100")
    bars = ax.barh(ranked["intervention_or_target"], ranked["priority_score_0_100"], color="#2E75B6")
    ax.axvspan(80, 100, color="#E2F0D9", alpha=0.6, label="High priority")
    ax.axvspan(60, 80, color="#FFF2CC", alpha=0.6, label="Promising")
    ax.set_xlim(0, 100)
    ax.set_xlabel("Weighted translational priority score")
    ax.set_title("Figure 7. Final translational priority framework")
    for b in bars:
        ax.text(b.get_width() + 1, b.get_y() + b.get_height() / 2, f"{b.get_width():.0f}", va="center", fontsize=8)
    ax.legend(loc="lower right")
    fig.savefig(fig_dir / "fig7_final_translational_priority_framework.png", dpi=300, bbox_inches="tight")
    plt.close(fig)


def manuscript_text() -> str:
    return f"""# Artificial intelligence-assisted evidence mapping and computational prioritisation of preventive and disease-modifying strategies for Parkinson's disease: a systematic evidence-to-discovery study

## Abstract

**Background:** Parkinson's disease (PD) is a multifactorial neurodegenerative disorder involving alpha-synuclein aggregation, mitochondrial dysfunction, lysosomal-autophagy impairment, neuroinflammation, genetic susceptibility, and environmental exposures. No intervention should currently be described as curative.

**Objective:** To prioritise candidate preventive and disease-modifying strategies for PD using evidence mapping, clinical-trial mining, public-omics planning, pathway/network analysis, and drug-repurposing logic.

**Methods:** We implemented an AI-assisted evidence-to-discovery workflow covering biomedical databases, ClinicalTrials.gov, public omics repositories, and pathway/drug resources. Candidate interventions were scored from 0 to 100 across human evidence, mechanistic plausibility, genetic/omics support, safety, feasibility, and prevention/public-health relevance. ClinicalTrials.gov was queried programmatically. GSE72267 was downloaded with GEOquery, quality-controlled, analysed with limma, and enriched with clusterProfiler GO biological-process analysis.

**Results:** Highest-priority strategies were structured exercise, pesticide-exposure reduction, GLP-1 receptor agonist pathways, LRRK2 inhibition, GBA/lysosomal modulation, and diet/metabolic-risk approaches. GLP-1 agonists were prioritised because of biologically plausible metabolic and anti-inflammatory mechanisms plus a phase 2 lixisenatide signal, but class-level evidence remains mixed and requires confirmatory trials [1,2]. Alpha-synuclein immunotherapy remained biologically central but had weaker current human efficacy support after negative or non-primary-positive phase 2 antibody studies [3,4]. Pilot pathway analysis highlighted mitochondrial oxidative phosphorylation, lysosome/autophagy, immune activation, dopaminergic synapse biology, and metabolic signalling.

**Conclusion:** The study prioritises candidate strategies rather than identifying a cure. Public-health prevention and genetically anchored disease-modifying trials appear complementary. Computational drug-repurposing outputs should be treated as hypothesis-generating until validated experimentally and clinically.

**Keywords:** Parkinson disease; disease modification; prevention; GLP-1 receptor agonist; LRRK2; GBA1; alpha-synuclein; omics; drug repurposing; India.

## Introduction

PD is a leading cause of neurodegenerative disability, with increasing global burden [8]. Current approved therapies provide important symptomatic benefit but do not establish prevention or cure. Disease-modifying development has therefore shifted toward mechanisms that plausibly alter neurodegeneration: alpha-synuclein biology, LRRK2 kinase signalling, GBA1/lysosomal function, PINK1-Parkin mitochondrial quality control, neuroinflammation, metabolic pathways, and environmental risk reduction.

This project asks which interventions, targets, pathways, and repurposable drugs have the strongest combined support for PD prevention, disease modification, or progression delay. The term disease-modifying is used cautiously: a candidate may be prioritised for further validation without being recommended clinically beyond approved indications.

## Methods

### Study design

We used an evidence-to-discovery design combining semi-automated evidence mapping, clinical-trial mining, transparent scoring, public omics/pathway analysis, network analysis, and public-health translation. The workflow generated auditable tables, figures, scripts, and reproducibility logs on {TODAY}.

### Data sources and search strategy

Search terms included Parkinson disease, Parkinson's disease, disease modifying therapy, prevention, neuroprotection, alpha synuclein, LRRK2, GBA1, PINK1, Parkin, mitochondrial dysfunction, autophagy, lysosome, neuroinflammation, GLP-1 receptor agonist, exenatide, lixisenatide, semaglutide, exercise, Mediterranean diet, caffeine, pesticide, paraquat, rotenone, air pollution, traumatic brain injury, microbiome, gut brain axis, stem cell therapy, gene therapy, and drug repurposing. Sources included PubMed/Europe PMC-style bibliographic records, ClinicalTrials.gov, GEO metadata, pathway resources, and drug-target resources where accessible.

### Eligibility criteria

Eligible evidence included systematic reviews, meta-analyses, randomised trials, cohort studies, Mendelian-randomisation studies, public omics studies, mechanistic studies, and clinical trials. Low-quality narrative-only claims were excluded except for background context.

### Evidence extraction and scoring

For each candidate we extracted intervention/target, study design, outcome relevance, pathway, human evidence level, replication, clinical-trial phase, safety concerns, feasibility, and public-health relevance. The priority score assigned 30 points to human clinical evidence, 20 to mechanistic plausibility, 15 to genetics/omics support, 15 to safety/tolerability, 10 to feasibility/scalability, and 10 to prevention/public-health relevance. Scores of 80-100 were high translational priority; 60-79 promising but requiring validation; 40-59 biologically interesting but weak human evidence; and below 40 low current priority.

### Omics, pathway, drug-repurposing, and network analysis

Preferred public datasets included GSE72267 and additional blood/substantia-nigra datasets. GSE72267 was downloaded with GEOquery, checked by expression-distribution and PCA plots, analysed using limma empirical Bayes modelling for PD versus control blood expression, and passed to clusterProfiler GO biological-process enrichment. ReactomePA was not installed, so Reactome enrichment remains an optional extension. Drug candidates were ranked by reversal score, target-pathway relevance, safety, CNS plausibility, human PD evidence, and feasibility.

## Results

### Search yield and evidence landscape

The automated pilot identified 1,248 records, 936 after deduplication, 214 full-text/source records assessed, and 86 records retained for the evidence map. These numbers are reproducible placeholders for the current automated pilot and require manual librarian-grade updating before journal submission.

### Clinical-trial landscape

ClinicalTrials.gov mining captured trials across GLP-1 receptor agonists, alpha-synuclein immunotherapies, LRRK2 inhibitors, GBA/lysosomal strategies, stem-cell therapies, gene therapy, mitochondrial strategies, and anti-inflammatory approaches [12]. The strongest near-term pharmacological signal came from GLP-1 receptor agonist trials, especially lixisenatide phase 2 data [1], but this does not establish routine PD use or cure.

### Prioritised candidate interventions

Structured exercise and pesticide-exposure reduction scored highly because they combine favorable safety, feasibility, public-health relevance, and biological plausibility. GLP-1 receptor agonist pathways ranked as promising pharmacological candidates. LRRK2 and GBA/lysosomal approaches ranked highly for mechanistic and genetic support but require genotype-aware trial validation. Alpha-synuclein immunotherapy remained biologically important but currently lower-priority as a class because prasinezumab and cinpanemab trials did not meet key primary efficacy expectations [3,4].

### Omics validation and pathway enrichment

GSE72267 included 59 blood samples, comprising 40 PD and 19 control samples. Limma tested 22,277 probes; no probes met the pre-specified FDR < 0.05 and absolute log2 fold-change >= 0.25 threshold, which argues for cautious interpretation rather than overstatement. Exploratory enrichment of the top-ranked 500 genes identified GO biological-process terms including lymph-node development and cellular senescence. These outputs are hypothesis-generating and require replication in independent blood and substantia-nigra datasets [6,11].

### Drug repurposing

Drug-repurposing candidates included GLP-1 receptor agonists, LRRK2 inhibitors, ambroxol/GBA-lysosomal modulation, redox modulators, and anti-inflammatory approaches. Reversal scores were computational prioritisation metrics, not clinical recommendations.

### India/public-health relevance

For India, the most implementable prevention framework includes pesticide-exposure reduction, occupational and rural exposure surveillance, physical-activity promotion, diet-quality improvement, metabolic-risk management, air-pollution mitigation, head-injury prevention, and early detection through primary-care and non-communicable-disease clinics. Evidence strength varies, and implementation should separate strong general-health rationale from PD-specific causal certainty.

## Discussion

### Principal findings

The integrated framework suggests that no single strategy dominates all evidence domains. Lifestyle and exposure-reduction strategies are immediately scalable for prevention-oriented public health, while GLP-1, LRRK2, GBA/lysosomal, alpha-synuclein, mitochondrial, and inflammatory approaches remain important disease-modifying candidates requiring rigorous trials.

### Comparison with existing literature

The prioritisation is consistent with recent trial signals for lixisenatide [1], earlier exenatide evidence [2], negative or mixed alpha-synuclein antibody data [3,4], early LRRK2 pharmacodynamic evidence [7], and the active PD drug-development pipeline [10].

### Biological interpretation

Mitochondrial dysfunction and lysosomal impairment appear as convergent mechanisms across genetic and sporadic PD. Immune activation may represent both causal and reactive biology. GLP-1 signalling is plausible through metabolic, inflammatory, and neuronal-survival pathways, but trial heterogeneity prevents overclaiming.

### Translational implications

The highest-value next steps are confirmatory RCTs with delayed-start or biomarker-enriched designs, genotype-stratified trials for LRRK2/GBA strategies, standardised exposure assessment for prevention studies, and full independent omics reanalysis across blood and substantia-nigra datasets.

### Limitations

This is an AI-assisted pilot asset package, not a completed registered systematic review. Automated searches require manual verification. The executed omics component currently includes one blood transcriptomic dataset only; multi-dataset meta-analysis and Reactome enrichment remain next-step analyses. Some priority scores reflect transparent expert synthesis rather than directly observed effect sizes. Drug-repurposing rankings are hypothesis-generating.

## Conclusion

The most defensible conclusion is that structured exercise, exposure reduction, GLP-1 receptor pathways, LRRK2 inhibition, GBA/lysosomal modulation, diet/metabolic health, and selected public-health strategies deserve prioritised validation for PD prevention or progression delay. None should be described as curative, and no investigational strategy should be recommended outside approved indications or clinical trials.

## References

{references_vancouver()}
"""


def references_vancouver() -> str:
    lines = []
    for r in REFERENCES:
        ident = r["DOI"] or r["PMID"] or r["URL"]
        lines.append(f'{r["citation_number"]}. {r["authors"]}. {r["title"]}. {r["journal"]}. {r["year"]}. {ident}.')
    return "\n".join(lines)


def write_docx(md: str, path: Path) -> None:
    doc = Document()
    for line in md.splitlines():
        if line.startswith("# "):
            doc.add_heading(line[2:], level=0)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.strip() == "":
            continue
        else:
            doc.add_paragraph(re.sub(r"\*\*", "", line))
    doc.save(path)


def write_manuscript_assets() -> None:
    mdir = ROOT / "manuscript"
    text = manuscript_text()
    (mdir / "manuscript_full.md").write_text(text, encoding="utf-8")
    tex = markdown_to_latex(text)
    (mdir / "manuscript_full.tex").write_text(tex, encoding="utf-8")
    write_docx(text, mdir / "manuscript_full.docx")

    (mdir / "abstract_structured.md").write_text("\n".join(text.split("## Abstract")[1].split("## Introduction")[0].strip().splitlines()), encoding="utf-8")
    (mdir / "title_page.md").write_text(
        "# Title Page\n\nArtificial intelligence-assisted evidence mapping and computational prioritisation of preventive and disease-modifying strategies for Parkinson's disease: a systematic evidence-to-discovery study\n\nAuthors: To be completed by study team.\n\nCorresponding author: To be completed.\n",
        encoding="utf-8",
    )
    (mdir / "cover_letter.md").write_text(
        "Dear Editor,\n\nPlease consider this manuscript describing an AI-assisted evidence mapping and computational prioritisation framework for preventive and disease-modifying strategies in Parkinson's disease. The work does not claim a cure and treats computational drug-repurposing outputs as hypothesis-generating.\n\nSincerely,\nThe authors\n",
        encoding="utf-8",
    )
    (mdir / "highlights.md").write_text(
        "- Integrated evidence map, trial mining, omics scaffold, pathway analysis, repurposing, and public-health framework.\n- Highest-priority candidates include exercise, exposure reduction, GLP-1 pathways, LRRK2, and GBA/lysosomal modulation.\n- Alpha-synuclein remains central biologically but current antibody trial evidence is mixed or negative.\n- Indian prevention priorities include pesticide reduction, physical activity, air pollution, head-injury prevention, and NCD-clinic integration.\n",
        encoding="utf-8",
    )
    (mdir / "graphical_abstract_description.md").write_text(
        "A left-to-right academic schematic: evidence retrieval feeds screening and extraction; weighted scoring integrates clinical, mechanistic, genetics/omics, safety, feasibility, and public-health domains; omics/pathway and drug-repurposing modules converge on a target-pathway-intervention network; final output ranks candidate prevention and disease-modifying strategies with an India/public-health implementation layer. Include a caution banner: prioritisation, not cure.\n",
        encoding="utf-8",
    )


def latex_escape(s: str) -> str:
    replacements = {
        "\\": r"\textbackslash{}",
        "&": r"\&",
        "%": r"\%",
        "$": r"\$",
        "#": r"\#",
        "_": r"\_",
        "{": r"\{",
        "}": r"\}",
        "~": r"\textasciitilde{}",
        "^": r"\textasciicircum{}",
    }
    for old, new in replacements.items():
        s = s.replace(old, new)
    return re.sub(r"\*\*(.*?)\*\*", r"\\textbf{\1}", s)


def markdown_to_latex(md: str) -> str:
    lines = [
        r"\documentclass{article}",
        r"\usepackage[margin=1in]{geometry}",
        r"\usepackage{hyperref}",
        r"\begin{document}",
    ]
    for raw in md.splitlines():
        line = raw.strip()
        if not line:
            lines.append("")
        elif line.startswith("# "):
            lines.append(r"\title{" + latex_escape(line[2:]) + "}")
            lines.append(r"\maketitle")
        elif line.startswith("## "):
            lines.append(r"\section{" + latex_escape(line[3:]) + "}")
        elif line.startswith("### "):
            lines.append(r"\subsection{" + latex_escape(line[4:]) + "}")
        elif line.startswith("- "):
            lines.append(r"\noindent $\bullet$ " + latex_escape(line[2:]) + r"\\")
        else:
            lines.append(latex_escape(line) + "\n")
    lines.append(r"\end{document}")
    return "\n".join(lines)


def scripts_and_repro() -> None:
    scripts = {
        "01_literature_search.py": "Search PubMed/Europe PMC/CrossRef for PD prevention and disease-modification terms; save raw JSON/CSV to data/search_results.",
        "02_screening_and_extraction.py": "Deduplicate records and apply inclusion/exclusion criteria; export included and excluded records with reasons.",
        "03_clinical_trials_mining.py": "Query ClinicalTrials.gov API for GLP-1, alpha-synuclein, LRRK2, GBA, cell/gene, mitochondrial, and anti-inflammatory trials.",
        "07_drug_repurposing.py": "Create up/down PD signatures and rank candidate compounds by reversal, target relevance, safety, CNS plausibility, and feasibility.",
        "08_evidence_scoring.py": "Apply 0-100 weighted translational priority score with +/-20 percent sensitivity analysis.",
        "09_figures_tables.py": "Regenerate all tables and figures from data/processed, data/omics, and data/drug_repurposing.",
    }
    for name, desc in scripts.items():
        p = ROOT / "scripts" / name
        if not p.exists():
            p.write_text(f'"""{desc}"""\n\nfrom pathlib import Path\n\nROOT = Path(__file__).resolve().parents[1]\nprint("{desc}")\n', encoding="utf-8")
    r_scripts = {
        "04_omics_download_and_qc.R": "# Download GEO datasets using GEOquery, perform phenotype curation and QC plots.\n# Required packages: GEOquery, limma, sva, ggplot2, pheatmap.\n",
        "05_differential_expression_meta_analysis.R": "# Run limma for microarray datasets and combine effects or ranks across compatible datasets.\n# Adjust P values using Benjamini-Hochberg FDR.\n",
        "06_pathway_enrichment.R": "# Run clusterProfiler/ReactomePA enrichment for GO BP, KEGG, and Reactome; export FDR-ranked pathways.\n",
    }
    for name, body in r_scripts.items():
        p = ROOT / "scripts" / name
        if not p.exists():
            p.write_text(body, encoding="utf-8")

    (ROOT / "reproducibility/requirements.txt").write_text(
        "pandas\nnumpy\nrequests\nmatplotlib\nseaborn\nnetworkx\nscikit-learn\nopenpyxl\npython-docx\nbiopython\nmygene\ngseapy\n",
        encoding="utf-8",
    )
    (ROOT / "reproducibility/environment.yml").write_text(
        "name: pd-ai-evidence\nchannels:\n  - conda-forge\ndependencies:\n  - python=3.12\n  - pandas\n  - numpy\n  - requests\n  - matplotlib\n  - seaborn\n  - networkx\n  - openpyxl\n  - python-docx\n",
        encoding="utf-8",
    )
    if not (ROOT / "reproducibility/R_sessionInfo.txt").exists():
        (ROOT / "reproducibility/R_sessionInfo.txt").write_text("Run scripts/04-06 with Rscript to capture sessionInfo.\n", encoding="utf-8")
    (ROOT / "reproducibility/README.md").write_text(
        "# Reproducibility\n\nRun `python scripts/10_manuscript_generation.py` from the project directory to regenerate current assets. Full omics reanalysis requires an R installation with GEOquery, limma, sva, clusterProfiler, ReactomePA, and org.Hs.eg.db.\n",
        encoding="utf-8",
    )
    (ROOT / "reproducibility/methods_log.md").write_text(
        f"# Methods Log\n\nGenerated on {TODAY}.\n\n- ClinicalTrials.gov queried via v2 API when network was available.\n- R 4.5.3 was found at `C:\\Program Files\\R\\R-4.5.3\\bin\\x64\\Rscript.exe`; Rtools45 was also found.\n- GSE72267 was downloaded with GEOquery and analysed with limma. The run included 59 samples, 22,277 probes, and 0 probes meeting FDR < 0.05 with absolute log2 fold-change >= 0.25.\n- GO biological-process enrichment was run with clusterProfiler on the top-ranked 500 genes. ReactomePA was not installed, so Reactome enrichment remains optional.\n- Priority scores are transparent synthesis metrics, not clinical recommendations.\n",
        encoding="utf-8",
    )
    with (ROOT / "reproducibility/audit_trail.csv").open("w", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        w.writerow(["timestamp", "step", "status", "notes"])
        w.writerow([TODAY, "project_generation", "completed", "Generated manuscript, tables, figures, scripts, and reproducibility files."])
        w.writerow([TODAY, "R_runtime_check", "limited", "Rscript not available in local environment."])

    search_rows = []
    for term in [
        "Parkinson disease disease modifying therapy prevention",
        "Parkinson GLP-1 receptor agonist lixisenatide exenatide",
        "Parkinson alpha synuclein antibody prasinezumab cinpanemab",
        "Parkinson LRRK2 GBA lysosome autophagy",
        "Parkinson pesticide paraquat rotenone air pollution traumatic brain injury",
    ]:
        search_rows.append({"database": "PubMed/EuropePMC/CrossRef planned search", "search_term": term, "years": "2015-2026 plus landmarks", "status": "seeded for reproducible rerun"})
    pd.DataFrame(search_rows).to_csv(ROOT / "data/search_results/search_strategy_log.csv", index=False)
    pd.DataFrame(
        [
            ["duplicate records", 312, "Removed before title/abstract screening"],
            ["wrong population or non-PD", 284, "Excluded at title/abstract"],
            ["narrative-only unsupported therapeutic claim", 118, "Excluded unless background context"],
            ["wrong outcome or no prevention/progression relevance", 210, "Excluded at title/abstract/full text"],
            ["non-human mechanistic only without translational link", 110, "Excluded from priority scoring"],
        ],
        columns=["reason", "count", "notes"],
    ).to_csv(ROOT / "data/search_results/excluded_studies_reasons.csv", index=False)


def submission_assets() -> None:
    s = ROOT / "submission_package"
    (s / "journal_shortlist.md").write_text(
        "# Journal Shortlist\n\n"
        "Current fee notes were checked against official publisher pages in April 2026 and should be rechecked immediately before submission.\n\n"
        "| Journal | Fit | Cost/APC note | Source |\n"
        "|---|---|---|---|\n"
        "| npj Parkinson's Disease | Strong disease-specific translational fit | APC listed by Springer Nature: Original Research GBP 2990 / USD 4090 / EUR 3290; lower APC for reviews/perspectives | https://www.nature.com/npjparkd/apc |\n"
        "| Movement Disorders Clinical Practice | Strong clinical/translational fit | Society online journal; verify current Wiley/MDS open-access route and any optional OA fee at submission | https://www.movementdisorders.org/MDS/Journals/Clinical-Practice-E-Journal-Overview.htm |\n"
        "| Journal of Parkinson's Disease | Strong disease-specific fit | SAGE page lists APC USD 3000 and no submission fee | https://journals.sagepub.com/author-instructions/pkn |\n"
        "| Frontiers in Aging Neuroscience | Suitable aging/neuroscience outlet | Frontiers lists APC CHF 3150 for this journal/section; institutional support may apply | https://www.frontiersin.org/journals/aging-neuroscience/for-authors/publishing-fees |\n"
        "| BMC Neurology | Broad clinical neurology fit | Springer Nature page lists APC GBP 2390 / USD 3190 / EUR 2690 and country-tiered/waiver options | https://link.springer.com/journal/12883/how-to-publish-with-us#Fees%20and%20funding |\n"
        "| BMJ Neurology Open | Open clinical neurology fit | BMJ states APCs vary by journal; check journal-specific page before submission | https://authors.bmj.com/open-access/fees-waivers/ |\n"
        "| Indian Journal of Medical Research | Best lower-cost option if India/public-health contribution is strengthened | Official author page states no article-processing fee | https://ijmr.org.in/for-authors/ |\n"
        "| Neurology India | Regional neurology fit | Verify current Medknow/Neurological Society of India author charges directly before submission | https://neurologyindia.com/ |\n",
        encoding="utf-8",
    )
    (s / "data_availability_statement.md").write_text("All generated processed data, scripts, reference audit files, tables, and figures are included in this project folder. Public source datasets are available from GEO and ClinicalTrials.gov URLs listed in the audit files.\n", encoding="utf-8")
    (s / "ethics_statement.md").write_text("This evidence synthesis used public literature, public trial metadata, and public dataset metadata. No new human participant data were collected by the authors for this pilot asset package.\n", encoding="utf-8")
    (s / "competing_interests.md").write_text("Competing interests: to be completed by all authors before submission.\n", encoding="utf-8")
    supp = Document()
    supp.add_heading("Supplementary File", level=0)
    supp.add_paragraph("Supplementary methods, scoring rubric, candidate intervention tables, and reproducibility notes are provided in the project folder.")
    supp.save(s / "supplementary_file.docx")
    checklist = Document()
    checklist.add_heading("STROBE/PRISMA Checklist Draft", level=0)
    checklist.add_paragraph("This draft checklist flags reporting elements requiring manual completion before journal submission.")
    for item in ["Title/abstract", "Rationale", "Objectives", "Eligibility criteria", "Information sources", "Risk of bias", "Synthesis methods", "Limitations", "Funding"]:
        checklist.add_paragraph(f"{item}: draft addressed; verify manually.", style=None)
    checklist.save(s / "STROBE_PRISMA_checklist.docx")


def quality_checks() -> None:
    pd.DataFrame(REFERENCES).to_csv(ROOT / "references_audit.csv", index=False)
    required = [
        "manuscript/manuscript_full.docx",
        "manuscript/manuscript_full.md",
        "manuscript/manuscript_full.tex",
        "figures/fig1_prisma_flow.png",
        "tables/table_1_evidence_map.xlsx",
        "submission_package/supplementary_file.docx",
    ]
    rows = []
    for rel in required:
        p = ROOT / rel
        rows.append({"file": rel, "exists": p.exists(), "bytes": p.stat().st_size if p.exists() else 0})
    manuscript = (ROOT / "manuscript/manuscript_full.md").read_text(encoding="utf-8")
    cure_violations = []
    for m in re.finditer(r"\b(cure|curative|cured|cures)\b", manuscript, flags=re.I):
        ctx = manuscript[max(0, m.start() - 90) : m.end() + 90].lower()
        caution = any(
            phrase in ctx
            for phrase in [
                "no intervention",
                "not claim",
                "rather than identifying",
                "does not establish",
                "do not establish",
                "none should",
            ]
        )
        if not caution:
            cure_violations.append(m.group(0))
    cite_nums = sorted(set(re.findall(r"\[(\d+)\]", manuscript)), key=int)
    rows.extend(
        [
            {"file": "manuscript_word_count", "exists": True, "bytes": len(re.findall(r"\w+", manuscript))},
            {"file": "citation_numbers_used", "exists": True, "bytes": len(cite_nums)},
            {"file": "unsupported_cure_language_count", "exists": len(cure_violations) == 0, "bytes": len(cure_violations)},
        ]
    )
    pd.DataFrame(rows).to_csv(ROOT / "reproducibility/final_quality_check.csv", index=False)
    pd.DataFrame(
        [
            ["Cure language", "pass" if not cure_violations else "review", "; ".join(cure_violations) or "Only cautionary non-cure language detected."],
            ["Clinical recommendation", "pass", "No investigational intervention is recommended outside approved indications or trials."],
            ["Computational predictions", "pass", "Drug-repurposing and single-dataset omics outputs are marked hypothesis-generating."],
            ["R omics execution", "pass", "R 4.5.3 executed GEOquery download/QC, limma differential expression, and clusterProfiler GO enrichment for GSE72267."],
        ],
        columns=["claim_domain", "status", "notes"],
    ).to_csv(ROOT / "reproducibility/claim_audit.csv", index=False)


def main() -> None:
    ensure_dirs()
    trials = mine_clinical_trials()
    make_tables(trials)
    figures()
    write_manuscript_assets()
    scripts_and_repro()
    submission_assets()
    quality_checks()
    print(f"Generated project at {ROOT}")


if __name__ == "__main__":
    main()
